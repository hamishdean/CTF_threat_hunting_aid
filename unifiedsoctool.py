# -*- coding: utf-8 -*-

import os
import sys
import json
import time
import math
import threading
import re
import datetime
import tkinter as tk
from tkinter import ttk, messagebox, filedialog, simpledialog
from tkinter import scrolledtext
from datetime import timedelta
from typing import List, Set

# --- Third Party Imports ---
try:
    from pypdf import PdfReader
    import pandas as pd
    from colorama import Fore, init, Style
    from openai import OpenAI
    from azure.identity import DefaultAzureCredential
    from azure.monitor.query import LogsQueryClient
    from azure.core.exceptions import HttpResponseError

    # Imports for Report Generator & Docx Reading
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError as e:
    print(f"Missing Library: {e}")
    print("Please install: pip install openai pypdf azure-identity azure-monitor-query pandas colorama tiktoken python-docx")
    print("Optional AI providers: pip install anthropic google-generativeai")

# Optional AI Provider imports
HAS_ANTHROPIC = False
try:
    import anthropic
    HAS_ANTHROPIC = True
except ImportError:
    pass

HAS_GEMINI = False
try:
    import google.generativeai as genai
    HAS_GEMINI = True
except ImportError:
    pass

# ==========================================
# 1. SHARED CONFIGURATION & UTILS
# ==========================================

MODEL_OPTIONS = ["gpt-4o", "gpt-4o-mini", "gpt-5.2", "gpt-5-mini"]
DEFAULT_MODEL = "gpt-4o"

PROVIDER_OPTIONS = ["OpenAI", "Gemini", "Claude"]
PROVIDER_MODELS = {
    "OpenAI": ["gpt-4o", "gpt-4o-mini", "gpt-5.2", "gpt-5-mini"],
    "Gemini": ["gemini-2.5-pro", "gemini-2.5-flash", "gemini-2.0-flash", "gemini-3.1-pro", "gemini-3.1-flash"],
    "Claude": ["claude-sonnet-4-6", "claude-opus-4-6", "claude-haiku-4-5-20251001"],
}
PROVIDER_DEFAULTS = {
    "OpenAI": "gpt-4o",
    "Gemini": "gemini-2.5-pro",
    "Claude": "claude-sonnet-4-6",
}
PROVIDER_ENV_KEYS = {
    "OpenAI": "OPENAI_API_KEY",
    "Gemini": "GEMINI_API_KEY",
    "Claude": "ANTHROPIC_API_KEY",
}

class SessionLogger:
    """Redirects print statements to a tkinter ScrolledText widget."""
    def __init__(self, text_widget):
        self.text_widget = text_widget
        self.ansi_escape = re.compile(r'\x1B(?:[@-Z\\-_]|\[[0-?]*[ -/]*[@-~])')

    def write(self, message):
        # Clean ANSI color codes for GUI display
        clean_msg = self.ansi_escape.sub('', message)

        if self.text_widget:
            self.text_widget.config(state="normal")
            self.text_widget.insert("end", clean_msg)
            self.text_widget.see("end")
            self.text_widget.config(state="disabled")

    def flush(self):
        pass

    def save_to_file(self, filename):
        try:
            with open(filename, "w", encoding="utf-8") as f:
                f.write(self.text_widget.get("1.0", "end"))
            return True
        except Exception as e:
            return False

class AzureSentinelFormatter:
    """Formats output to look like Azure Log Analytics / Sentinel results."""
    @staticmethod
    def format_log(finding: dict) -> str:
        formatted = {
            "TimeGenerated": finding.get("timestamp", "2025-12-01T00:00:00Z"),
            "AlertName": finding.get("title", "Unknown Alert"),
            "FlagAnswer": finding.get("flag_answer", ""),
            "Description": finding.get("description", ""),
            "Severity": finding.get("severity", "Medium"),
            "Evidence": finding.get("evidence", "")
        }
        return json.dumps(formatted, indent=2)

def extract_text_from_file(filepath: str) -> List[str]:
    """Reads PDF, DOCX, or TXT files and returns a list of page strings (or chunks)."""
    text_chunks = []
    try:
        ext = filepath.lower()
        if ext.endswith('.pdf'):
            reader = PdfReader(filepath)
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_chunks.append(text)
        elif ext.endswith('.docx'):
            doc = Document(filepath)
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            text_chunks.append('\n'.join(full_text))
        else:
            with open(filepath, 'r', encoding='utf-8', errors='replace') as f:
                content = f.read()
                chunk_size = 4000
                for i in range(0, len(content), chunk_size):
                    text_chunks.append(content[i:i+chunk_size])
    except Exception as e:
        print(f"Error reading {filepath}: {e}")
    return text_chunks

# ==========================================
# 2. PROMPTS & AI UTILS
# ==========================================

# UPDATED: Matched strictly to EXECUTOR.py reference for better performance
SYSTEM_PROMPT_KQL_GENERATOR = """
You are an expert Azure KQL (Kusto Query Language) engineer.
Translate the user's request into a SHORT, SIMPLE, VALID KQL query.

### CRITICAL: KEEP QUERIES SHORT
- Write the SIMPLEST query that answers the question. Prefer 1-3 pipes max.
- NEVER use subqueries, let statements, joins, or unions unless the user explicitly asks.
- NEVER use functions like make_list, mv-expand, arg_max, bag_unpack, todynamic unless essential.
- If unsure which table to use, use `search "keyword" | take 100` as a safe fallback.
- ALWAYS end with `| take 500` (or less) to limit results.

### TABLES
You can query ANY table in the Azure Log Analytics workspace. Use whatever table the user asks for.
If the user mentions a specific table name, use it directly.
If the user describes what they want but doesn't name a table, pick the best table from the common ones below or use `search`.

Common tables, what they contain, when to use them, and their typical columns (not exhaustive - tables may have additional columns):

- DeviceProcessEvents: Process creation and execution events on endpoints. Use for investigating suspicious commands, script execution, malware processes, LOLBins, and process trees. Columns: TimeGenerated, AccountName, ActionType, DeviceName, InitiatingProcessCommandLine, ProcessCommandLine, FileName, FolderPath
- DeviceNetworkEvents: Outbound/inbound network connections from devices. Use for C2 beaconing, data exfiltration, lateral movement, and suspicious remote connections. Columns: TimeGenerated, ActionType, DeviceName, RemoteIP, RemotePort, RemoteUrl, LocalIP
- DeviceLogonEvents: Authentication events on devices. Use for brute force detection, pass-the-hash, RDP sessions, and unusual logon patterns. Columns: TimeGenerated, AccountName, DeviceName, ActionType, RemoteIP, RemoteDeviceName, LogonType
- DeviceFileEvents: File creation, modification, and deletion on endpoints. Use for malware drops, file staging, ransomware activity, and suspicious file operations. Columns: TimeGenerated, DeviceName, FileName, FolderPath, ActionType, SHA256, InitiatingProcessFileName
- DeviceRegistryEvents: Windows registry changes on endpoints. Use for persistence mechanisms, autorun keys, config tampering, and malware registry modifications. Columns: TimeGenerated, DeviceName, ActionType, RegistryKey, RegistryValueName, RegistryValueData
- DeviceImageLoadEvents: DLL and driver load events on endpoints. Use for DLL sideloading, DLL injection, unsigned module loading, and suspicious driver loads. Columns: TimeGenerated, DeviceName, FileName, FolderPath, SHA256, InitiatingProcessFileName
- DeviceEvents: Miscellaneous device events including USB, AMSI, exploit guard, and PowerShell logging. Use for USB device connections, exploit mitigation triggers, and advanced threat events. Columns: TimeGenerated, DeviceName, ActionType, AdditionalFields
- EmailEvents: Email delivery and filtering events. Use for phishing investigations, spam analysis, malicious email delivery or blocking, and email-based attack chains. Columns: TimeGenerated, SenderFromAddress, RecipientEmailAddress, Subject, DeliveryAction, ThreatTypes
- EmailAttachmentInfo: Email attachment metadata. Use for investigating malicious attachments, suspicious file types in emails, and attachment-based threats. Columns: TimeGenerated, FileName, FileType, SHA256, SenderFromAddress, RecipientEmailAddress
- EmailUrlInfo: URLs found in email messages. Use for phishing link analysis, malicious URL detection, and URL-based email threats. Columns: TimeGenerated, Url, UrlDomain, NetworkMessageId
- IdentityLogonEvents: Identity-level sign-in events (Active Directory and cloud). Use for failed logon analysis, unusual authentication patterns, and identity-based attacks. Columns: TimeGenerated, AccountUpn, ActionType, LogonType, Application, IPAddress
- IdentityQueryEvents: LDAP and directory query events. Use for reconnaissance detection, enumeration of users/groups/computers, and suspicious directory queries. Columns: TimeGenerated, AccountUpn, ActionType, QueryType, QueryTarget
- IdentityDirectoryEvents: Active Directory changes. Use for account modifications, group membership changes, password resets, and privilege escalation via directory changes. Columns: TimeGenerated, AccountUpn, ActionType, TargetAccountUpn, Application
- CloudAppEvents: Cloud application activity (SaaS apps). Use for suspicious cloud actions, shadow IT detection, OAuth app abuse, and cloud-based data access. Columns: TimeGenerated, AccountId, ActionType, Application, IPAddress, UserAgent
- AlertEvidence: Evidence entities linked to security alerts. Use for correlating IPs, files, users, and accounts associated with triggered alerts. Columns: TimeGenerated, EntityType, EvidenceRole, RemoteIP, FileName, AccountName
- SigninLogs: Azure AD interactive sign-in events. Use for MFA failures, impossible travel detection, conditional access evaluation, and user sign-in anomalies. Columns: TimeGenerated, UserPrincipalName, IPAddress, AppDisplayName, ResultType, ResultDescription, Location, ConditionalAccessStatus
- AADNonInteractiveUserSignInLogs: Service and app token sign-in events. Use for compromised service principals, token abuse, background app authentication, and non-interactive access patterns. Columns: TimeGenerated, UserPrincipalName, IPAddress, AppDisplayName, ResultType
- AuditLogs: Azure AD configuration and administrative changes. Use for role assignments, app registrations, policy changes, consent grants, and directory configuration modifications. Columns: TimeGenerated, OperationName, Result, InitiatedBy, TargetResources, Category
- AzureActivity: Azure resource management operations. Use for unauthorized resource changes, privilege escalation via Azure RBAC, resource creation/deletion, and subscription-level activity. Columns: TimeGenerated, Caller, OperationNameValue, ActivityStatusValue, ResourceGroup, SubscriptionId
- SecurityAlert: Security alerts from Microsoft Defender and Sentinel analytics rules. Use for alert triage, attack detection, MITRE ATT&CK mapping, and initial incident investigation. Columns: TimeGenerated, AlertName, AlertSeverity, Description, CompromisedEntity, Tactics, Techniques, ProviderName
- SecurityIncident: Aggregated security incidents from Sentinel. Use for incident management, viewing linked alerts, tracking incident status, and high-level investigation overview. Columns: TimeGenerated, Title, Severity, Status, Owner, IncidentNumber
- SecurityEvent: Windows Security Event Log (legacy agent). Use for classic event ID analysis (4624/4625 logons, 4688 process creation, 4672 special privileges), Windows security auditing. Columns: TimeGenerated, EventID, Account, Computer, Activity, IpAddress
- Syslog: Linux and network device syslog messages. Use for Linux authentication logs, service crashes, firewall logs, and network device monitoring. Columns: TimeGenerated, Computer, Facility, SeverityLevel, SyslogMessage, ProcessName
- CommonSecurityLog: CEF-format logs from security appliances. Use for firewall logs, IDS/IPS alerts, proxy logs, and third-party security device events. Columns: TimeGenerated, DeviceVendor, DeviceProduct, Activity, SourceIP, DestinationIP, DestinationPort
- Event: Windows Event Log (non-security, legacy agent). Use for application errors, system events, and non-security Windows logs. Columns: TimeGenerated, EventID, Computer, Source, RenderedDescription, UserName
- Heartbeat: Agent health and connectivity status. Use for confirming which machines are reporting, agent uptime, and connectivity monitoring. Columns: TimeGenerated, Computer, OSType, Version, ComputerIP
- ThreatIntelligenceIndicator: Threat intelligence IOCs ingested into Sentinel. Use for matching known malicious IPs, domains, file hashes, and URLs against your environment. Columns: TimeGenerated, ThreatType, IndicatorId, DomainName, Url, NetworkIP
- BehaviorAnalytics: UEBA (User and Entity Behavior Analytics) anomalies. Use for detecting unusual user behavior, insider threats, and anomalous access patterns. Columns: TimeGenerated, UserPrincipalName, ActionType, ActivityInsights, InvestigationPriority
- Anomalies: ML-detected anomalies from Sentinel analytics. Use for identifying anomalous patterns with MITRE ATT&CK tactics and techniques mapping. Columns: TimeGenerated, AnomalyName, Description, Score, Tactics, Techniques
- W3CIISLog: IIS web server access logs. Use for web application attacks, suspicious HTTP requests, directory traversal, SQL injection attempts, and web scanning detection. Columns: TimeGenerated, sSiteName, sIP, csMethod, csUriStem, scStatus, cIP
- AzureDiagnostics: Azure resource diagnostic logs. Use for resource-level operations, errors, and detailed diagnostics from Azure services (Key Vault, SQL, App Gateway, etc.). Columns: TimeGenerated, ResourceType, OperationName, ResultType, Category

If the user asks for a table not listed above, use it anyway - any valid table name in the workspace will work.
You can also use `search in (Table1, Table2) "keyword"` to search across specific tables.
Use `union Table1, Table2` only if the user explicitly asks to combine tables.

### SYNTAX RULES
1. Do NOT use wildcards (*) in where clauses. Use `startswith`, `contains`, or `has` operators.
2. String literals must be quoted: `"value"` or `'value'`.
3. Time filters: `| where TimeGenerated > ago(7d)` or `| where TimeGenerated between (datetime(...) .. datetime(...))`.
4. Default time range: 365 days (`ago(365d)`) unless the user specifies otherwise. You can use up to 1 year.

### OUTPUT FORMAT
Return a SINGLE JSON object:
{
    "table_name": "DeviceProcessEvents",
    "kql_query": "DeviceProcessEvents | where DeviceName startswith 'ch-' | take 100",
    "rationale": "Searching for devices with prefix ch-",
    "fields": "TimeGenerated, DeviceName, ProcessCommandLine, AccountName",
    "parameters": { "time_range_hours": 8760 }
}

CRITICAL: "fields" must be a single comma-separated string, NOT individual characters.
"""

FORMATTING_INSTRUCTIONS = """
You MUST return valid JSON. Format:
{"findings": [{"title": "...", "description": "...", "flag_answer": "the exact flag or artifact value", "confidence": "High", "log_lines": [], "notes": ""}]}
- "title": A short name for the finding.
- "flag_answer": The EXACT value that answers the question (e.g., an IP address, username, filename, hash, flag string, command, URL). This is the most important field.
- "description": Brief explanation of why this is the answer.
- "log_lines": Relevant raw log lines as evidence.
If nothing found, return: {"findings": []}
"""

SYSTEM_PROMPT_THREAT_HUNT = {
    "role": "system",
    "content": f"""You are an expert CTF Flag Hunter and SOC Analyst AI.
Your PRIMARY goal is to find the specific FLAG ANSWER - the exact artifact, value, string, IP, username, hash, command, or evidence that answers the investigation question or CTF hint.

Do NOT just describe threats generically. Instead:
- Extract the EXACT value (IP address, username, hostname, file hash, command line, flag string, etc.)
- Put that exact value in the "flag_answer" field.
- If a flag looks like flag{{...}} or FLAG{{...}}, extract it exactly.
- If the answer is a specific artifact (e.g., "192.168.1.100", "malware.exe", "admin@corp.com"), extract it exactly.

{FORMATTING_INSTRUCTIONS}"""
}

THREAT_HUNT_PROMPTS = {
    "DeviceProcessEvents": "Find the exact suspicious process, command line, filename, or script. Extract the specific executable name, full command, or file path as the flag_answer.",
    "DeviceLogonEvents": "Find the exact account, IP address, or device involved in suspicious logons. Extract the specific username, source IP, or device name as the flag_answer.",
    "DeviceNetworkEvents": "Find the exact remote IP, port, or connection details. Extract the specific IP address, domain, or port as the flag_answer.",
    "DeviceFileEvents": "Find the exact file name, path, or hash involved. Extract the specific filename, folder path, or SHA256 hash as the flag_answer.",
    "DeviceRegistryEvents": "Find the exact registry key, value name, or data modified. Extract the specific registry path or value as the flag_answer.",
    "DeviceImageLoadEvents": "Find the exact DLL or image loaded. Extract the specific filename, path, or hash as the flag_answer.",
    "DeviceEvents": "Find the exact event details. Extract the specific action, device, or artifact as the flag_answer.",
    "EmailEvents": "Find the exact sender, recipient, subject, or threat. Extract the specific email address, subject line, or threat type as the flag_answer.",
    "EmailAttachmentInfo": "Find the exact attachment filename, type, or hash. Extract the specific value as the flag_answer.",
    "EmailUrlInfo": "Find the exact URL or domain in the email. Extract the specific URL or domain as the flag_answer.",
    "IdentityLogonEvents": "Find the exact identity, logon type, or application. Extract the specific UPN, IP, or app as the flag_answer.",
    "IdentityQueryEvents": "Find the exact query target or type. Extract the specific query details as the flag_answer.",
    "IdentityDirectoryEvents": "Find the exact directory change or target account. Extract the specific value as the flag_answer.",
    "CloudAppEvents": "Find the exact cloud app action, account, or IP. Extract the specific value as the flag_answer.",
    "AlertEvidence": "Find the exact evidence entity - IP, file, account. Extract the specific value as the flag_answer.",
    "SigninLogs": "Find the exact user, IP address, or application involved in suspicious sign-ins. Extract the specific UPN, IP, app name, or error code as the flag_answer.",
    "AADNonInteractiveUserSignInLogs": "Find the exact service principal or app sign-in. Extract the specific UPN, IP, or app as the flag_answer.",
    "AuditLogs": "Find the exact operation, initiator, or target. Extract the specific value as the flag_answer.",
    "AzureActivity": "Find the exact caller, operation, or resource involved. Extract the specific identity, action, or resource name as the flag_answer.",
    "SecurityAlert": "Find the exact alert details - the compromised entity, attacker IP, or specific artifact. Extract the specific value as the flag_answer.",
    "SecurityIncident": "Find the exact incident details - title, affected entities, or linked alerts. Extract the specific value as the flag_answer.",
    "SecurityEvent": "Find the exact Windows security event details - EventID, account, or activity. Extract the specific value as the flag_answer.",
    "Syslog": "Find the exact syslog message, process, or facility. Extract the specific value as the flag_answer.",
    "CommonSecurityLog": "Find the exact source/destination IP, port, or activity from the security device. Extract the specific value as the flag_answer.",
    "ThreatIntelligenceIndicator": "Find the exact IOC - domain, URL, IP, or hash. Extract the specific indicator value as the flag_answer.",
    "BehaviorAnalytics": "Find the exact anomalous behavior, user, or activity insight. Extract the specific value as the flag_answer.",
    "Anomalies": "Find the exact anomaly details and score. Extract the specific anomaly name or artifact as the flag_answer.",
}

# Generic fallback for any table not listed above
THREAT_HUNT_PROMPT_DEFAULT = "Analyze the log records carefully. Find the exact artifact, value, or answer relevant to the investigation. Extract specific values (IPs, usernames, hashes, commands, flags, etc.) as the flag_answer."

SYSTEM_PROMPT_INCIDENT_REPORT = """
You are a Senior Security Operations Center (SOC) Analyst writing a formal incident report.

You will receive data from multiple sources:
- Verified flag answers and findings from the analyst's investigation
- CTF hints and AI-generated clues
- A Flag Bank narrative summarizing the current incident understanding
- SOC Agent queries (KQL) and their results from Azure Log Analytics
- Optionally, uploaded documents (logs, reports, etc.)

Your task:
1. Synthesize ALL provided data into a coherent, professional incident report.
2. Include the EXACT flag answers/artifacts in the report (do not paraphrase them).
3. Reference the KQL queries and evidence used to find each answer.
4. Build a chronological timeline from timestamps in the data.
5. Fill in the report template provided by the analyst.
6. If data is missing for a template section, note it as "Not enough data available".
"""

def ai_chat_completion(provider, api_key, model, messages, json_mode=False, max_tokens=4096, temperature=None):
    """Unified AI completion wrapper supporting OpenAI, Gemini, and Claude.

    Args:
        provider: "OpenAI", "Gemini", or "Claude"
        api_key: API key string for the chosen provider
        model: Model name string
        messages: List of {"role": ..., "content": ...} dicts (OpenAI format)
        json_mode: If True, request JSON output
        max_tokens: Maximum tokens in response
        temperature: Optional temperature override

    Returns:
        Response content as a string
    """
    if provider == "OpenAI":
        client = OpenAI(api_key=api_key)
        kwargs = {"model": model, "messages": messages}
        if json_mode:
            kwargs["response_format"] = {"type": "json_object"}
        if max_tokens:
            # Newer OpenAI models (o-series, gpt-5) require max_completion_tokens
            needs_completion_tokens = model.startswith(("o1", "o3", "o4", "gpt-5"))
            token_param = "max_completion_tokens" if needs_completion_tokens else "max_tokens"
            kwargs[token_param] = max_tokens
        if temperature is not None:
            kwargs["temperature"] = temperature
        response = client.chat.completions.create(**kwargs)
        return response.choices[0].message.content

    elif provider == "Claude":
        if not HAS_ANTHROPIC:
            raise ImportError("anthropic package not installed. Run: pip install anthropic")
        client = anthropic.Anthropic(api_key=api_key)
        # Extract system message from messages list
        system_text = ""
        user_messages = []
        for msg in messages:
            if msg["role"] == "system":
                system_text += msg["content"] + "\n"
            else:
                user_messages.append(msg)
        if json_mode:
            system_text += "\nYou MUST return valid JSON only. No extra text outside the JSON object."
        if not user_messages:
            user_messages = [{"role": "user", "content": "Please proceed."}]
        kwargs = {
            "model": model,
            "max_tokens": max_tokens or 4096,
            "messages": user_messages,
        }
        if system_text.strip():
            kwargs["system"] = system_text.strip()
        if temperature is not None:
            kwargs["temperature"] = temperature
        response = client.messages.create(**kwargs)
        return response.content[0].text

    elif provider == "Gemini":
        if not HAS_GEMINI:
            raise ImportError("google-generativeai package not installed. Run: pip install google-generativeai")
        genai.configure(api_key=api_key)
        gen_config = {}
        if json_mode:
            gen_config["response_mime_type"] = "application/json"
        if temperature is not None:
            gen_config["temperature"] = temperature
        if max_tokens:
            gen_config["max_output_tokens"] = max_tokens
        # Convert messages to Gemini format
        system_text = ""
        contents = []
        for msg in messages:
            if msg["role"] == "system":
                system_text += msg["content"] + "\n"
            elif msg["role"] == "assistant":
                contents.append({"role": "model", "parts": [msg["content"]]})
            else:
                contents.append({"role": "user", "parts": [msg["content"]]})
        # Prepend system text to the first user message if present
        if system_text.strip() and contents:
            first_content = contents[0]["parts"][0]
            contents[0]["parts"][0] = system_text.strip() + "\n\n" + first_content
        model_obj = genai.GenerativeModel(model, generation_config=gen_config if gen_config else None)
        response = model_obj.generate_content(contents)
        return response.text

    else:
        raise ValueError(f"Unsupported AI provider: {provider}")


def sanitize_kql(kql):
    """Clean up AI-generated KQL to fix common problems."""
    if not kql:
        return kql
    # Strip markdown code fences the AI sometimes wraps queries in
    kql = kql.strip()
    if kql.startswith("```"):
        kql = kql.split("\n", 1)[-1] if "\n" in kql else kql[3:]
    if kql.endswith("```"):
        kql = kql[:-3]
    kql = kql.strip()

    # Ensure a take limit exists to prevent runaway queries
    kql_lower = kql.lower()
    if "| take " not in kql_lower and "| limit " not in kql_lower and "| count" not in kql_lower:
        kql = kql.rstrip().rstrip(";") + "\n| take 500"

    return kql

def clean_json_string(json_str):
    if not json_str: return "{}"
    start = json_str.find('{')
    end = json_str.rfind('}')
    if start != -1 and end != -1 and end > start:
        return json_str[start : end + 1]
    return json_str.replace("```json", "").replace("```", "").strip()

def get_query_context(provider, api_key, user_input, model, history=None, hints=None, active_focus="General/All", incident_context=""):
    # Keep context concise to avoid confusing the KQL generator
    hist_str = ""
    if history:
        # Only include the last 2 queries, KQL only (no bloated context)
        hist_str = "Previous queries:\n" + "\n".join([f"- {h['kql_query']}" for h in history[-2:]])

    hints_str = ""
    if hints:
        # Only include short hint text, not full clue analysis
        hint_lines = [h['hint'][:100] for h in hints[:5]]
        hints_str = f"Context clues ({active_focus}): " + "; ".join(hint_lines)

    inc_str = ""
    if incident_context:
        # Truncate incident context to avoid overwhelming the KQL prompt
        inc_str = incident_context[:500]

    user_msg = f"User Request: {user_input}"
    if hints_str:
        user_msg = f"{hints_str}\n{user_msg}"
    if inc_str:
        user_msg = f"{inc_str}\n{user_msg}"
    if hist_str:
        user_msg = f"{hist_str}\n{user_msg}"

    messages = [
        {"role": "system", "content": SYSTEM_PROMPT_KQL_GENERATOR},
        {"role": "user", "content": user_msg}
    ]
    try:
        content = ai_chat_completion(provider, api_key, model, messages, json_mode=True)
        content = clean_json_string(content)
        query_context = json.loads(content)

        # UPDATED: Clean parameters like EXECUTOR.py
        if "kql_query" not in query_context:
            query_context["kql_query"] = ""
        else:
            query_context["kql_query"] = sanitize_kql(query_context["kql_query"])
        if "parameters" not in query_context:
            query_context["parameters"] = {}

        params = query_context.get("parameters", {})
        try:
            # Flatten time range
            raw_hours = params.get("time_range_hours", 8760)
            query_context["time_range_hours"] = int(str(raw_hours).lower().replace("hours","").replace("hour","").strip())
        except (ValueError, AttributeError):
            query_context["time_range_hours"] = 8760

        # Flatten specific keys
        query_context["device_name"] = params.get("device_name", "")
        query_context["caller"] = params.get("caller", "")
        query_context["user_principal_name"] = params.get("user_principal_name", "")

        return query_context

    except Exception as e:
        return {"table_name": "Error", "kql_query": "", "error": str(e)}

def execute_kql(law_client, workspace_id, kql, hours=8760):
    try:
        # Increase limit for manual/custom queries like EXECUTOR.py suggests
        if hours > 8800: hours = 8800

        response = law_client.query_workspace(
            workspace_id=workspace_id,
            query=kql,
            timespan=timedelta(hours=hours)
        )
        if hasattr(response, 'tables'):
            tables = response.tables
        elif hasattr(response, 'partial_data'):
            tables = response.partial_data
        else:
            tables = []

        if not tables:
            return []

        table = tables[0]
        columns = [col.name if hasattr(col, 'name') else str(col) for col in table.columns]
        results = [{col: val for col, val in zip(columns, row)} for row in table.rows]
        return results
    except Exception as e:
        raise e

def hunt_on_records(provider, api_key, records, table_name, model, hints=None, found_flags=None, active_focus="General/All"):
    try:
        df = pd.DataFrame(records)
        csv_data = df.to_csv(index=False)
    except Exception:
        csv_data = str(records)

    context_str = ""
    if hints:
        context_str += f"CTF HINTS (FOCUS: {active_focus}):\n" + "\n".join([f"- [{h.get('id', 'General')}] {h['hint']}" for h in hints]) + "\n"
    if found_flags:
        context_str += "ALREADY FOUND FLAGS (Ignore these):\n" + json.dumps(list(found_flags)) + "\n"

    specific_instructions = THREAT_HUNT_PROMPTS.get(table_name, THREAT_HUNT_PROMPT_DEFAULT)

    prompt = f"""{context_str}
Instructions: {specific_instructions}
Analyze these {table_name} logs. Find the EXACT flag answer or specific artifact value that solves the question.
Do NOT just say "suspicious activity detected". Instead, extract the exact value (IP, username, hash, command, flag string, etc.) and put it in "flag_answer".
LOG DATA:
{csv_data}"""

    try:
        messages = [SYSTEM_PROMPT_THREAT_HUNT, {"role": "user", "content": prompt}]
        content = ai_chat_completion(provider, api_key, model, messages, json_mode=True)
        content = clean_json_string(content)
        return json.loads(content).get("findings", [])
    except Exception as e:
        raise e

# ==========================================
# 3. REPORT GENERATOR LOGIC
# ==========================================

class ThreatHuntReporterTab:
    """Encapsulates the HuntLogApp logic within a Frame."""
    def __init__(self, parent_frame):
        self.parent = parent_frame
        self.entries = []
        self.current_images = []
        self._init_ui()

    def _init_ui(self):
        toolbar = ttk.Frame(self.parent, padding=5)
        toolbar.pack(fill=tk.X)
        ttk.Label(toolbar, text="HuntLog Reporter Module", font=("Segoe UI", 9, "italic")).pack(side=tk.RIGHT)

        self.notebook = ttk.Notebook(self.parent)
        self.notebook.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

        self.input_tab = ttk.Frame(self.notebook, padding="15")
        self.help_tab = ttk.Frame(self.notebook, padding="15")

        self.notebook.add(self.input_tab, text="Entry Form")
        self.notebook.add(self.help_tab, text="Help & Instructions")

        self._init_input_tab()
        self._init_help_tab()

    def _init_input_tab(self):
        # Scrollable wrapper for the entire input tab
        canvas = tk.Canvas(self.input_tab, highlightthickness=0)
        v_scroll = ttk.Scrollbar(self.input_tab, orient="vertical", command=canvas.yview)
        scroll_frame = ttk.Frame(canvas)

        scroll_frame.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
        canvas.create_window((0, 0), window=scroll_frame, anchor="nw")
        canvas.configure(yscrollcommand=v_scroll.set)

        canvas.pack(side="left", fill="both", expand=True)
        v_scroll.pack(side="right", fill="y")

        # Enable mousewheel scrolling
        def _on_mousewheel(event):
            canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")
        def _on_linux_scroll_up(event):
            canvas.yview_scroll(-3, "units")
        def _on_linux_scroll_down(event):
            canvas.yview_scroll(3, "units")

        canvas.bind_all("<MouseWheel>", _on_mousewheel)
        canvas.bind_all("<Button-4>", _on_linux_scroll_up)
        canvas.bind_all("<Button-5>", _on_linux_scroll_down)

        # Resize inner frame width to match canvas
        def _on_canvas_configure(event):
            canvas.itemconfig(canvas.find_all()[0], width=event.width)
        canvas.bind("<Configure>", _on_canvas_configure)

        # --- All content goes into scroll_frame instead of self.input_tab ---
        meta_frame = ttk.LabelFrame(scroll_frame, text="Report Details", padding="10")
        meta_frame.pack(fill=tk.X, pady=(0, 10))

        ttk.Label(meta_frame, text="Author Name:").grid(row=0, column=0, sticky=tk.W, padx=5)
        self.author_entry = ttk.Entry(meta_frame, width=30)
        self.author_entry.grid(row=0, column=1, sticky=tk.W, padx=5)

        ttk.Label(meta_frame, text="Event/Engagement:").grid(row=0, column=2, sticky=tk.W, padx=5)
        self.event_entry = ttk.Entry(meta_frame, width=30)
        self.event_entry.grid(row=0, column=3, sticky=tk.W, padx=5)

        input_frame = ttk.LabelFrame(scroll_frame, text="Add New Finding / Challenge", padding="10")
        input_frame.pack(fill=tk.X, pady=(0, 10))
        input_frame.columnconfigure(1, weight=1)

        ttk.Label(input_frame, text="Challenge/Title:").grid(row=0, column=0, sticky=tk.NW, pady=5)
        self.title_entry = ttk.Entry(input_frame)
        self.title_entry.grid(row=0, column=1, sticky=tk.EW, padx=5, pady=5)

        ttk.Label(input_frame, text="Category:").grid(row=0, column=2, sticky=tk.NW, pady=5)
        self.cat_combo = ttk.Combobox(input_frame, values=["Network Forensics", "Endpoint Security", "Malware Analysis", "OSINT", "Log Analysis", "Web Exploitation", "Cryptography", "Cloud Security", "Other"])
        self.cat_combo.current(0)
        self.cat_combo.grid(row=0, column=3, sticky=tk.EW, padx=5, pady=5)

        ttk.Label(input_frame, text="Description:").grid(row=1, column=0, sticky=tk.NW, pady=5)
        self.desc_text_frame, self.desc_text = self.create_scrollable_text(input_frame, height=3)
        self.desc_text_frame.grid(row=1, column=1, columnspan=3, sticky=tk.EW, padx=5, pady=5)

        ttk.Label(input_frame, text="Methodology:").grid(row=2, column=0, sticky=tk.NW, pady=5)
        self.method_text_frame, self.method_text = self.create_scrollable_text(input_frame, height=4)
        self.method_text_frame.grid(row=2, column=1, columnspan=3, sticky=tk.EW, padx=5, pady=5)

        ttk.Label(input_frame, text="KQL / Query:").grid(row=3, column=0, sticky=tk.NW, pady=5)
        self.kql_text_frame, self.kql_text = self.create_scrollable_text(input_frame, height=3, font=("Courier New", 10))
        self.kql_text_frame.grid(row=3, column=1, columnspan=3, sticky=tk.EW, padx=5, pady=5)

        ttk.Label(input_frame, text="Flag / Artifact:").grid(row=4, column=0, sticky=tk.W, pady=5)
        self.flag_entry = ttk.Entry(input_frame)
        self.flag_entry.grid(row=4, column=1, columnspan=3, sticky=tk.EW, padx=5, pady=5)

        ttk.Label(input_frame, text="Screenshots:").grid(row=5, column=0, sticky=tk.NW, pady=5)
        img_control_frame = ttk.Frame(input_frame)
        img_control_frame.grid(row=5, column=1, columnspan=3, sticky=tk.EW, padx=5, pady=5)

        btn_box = ttk.Frame(img_control_frame)
        btn_box.pack(side=tk.LEFT, fill=tk.Y, padx=(0, 10))
        self.add_img_btn = ttk.Button(btn_box, text="Add Image + Caption", command=self.add_image_with_caption)
        self.add_img_btn.pack(fill=tk.X, pady=(0, 5))
        self.remove_img_btn = ttk.Button(btn_box, text="Remove Selected", command=self.remove_selected_image)
        self.remove_img_btn.pack(fill=tk.X)

        self.img_listbox = tk.Listbox(img_control_frame, height=4, width=60, font=("Segoe UI", 9))
        self.img_listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        btn_frame = ttk.Frame(input_frame)
        btn_frame.grid(row=6, column=0, columnspan=4, pady=10)
        ttk.Button(btn_frame, text="Add Entry to Report", command=self.add_entry).pack(side=tk.LEFT, padx=10)
        ttk.Button(btn_frame, text="Clear Form", command=self.clear_form).pack(side=tk.LEFT, padx=10)

        list_frame = ttk.LabelFrame(scroll_frame, text="Entries Queued for Export", padding="10")
        list_frame.pack(fill=tk.X, pady=(0, 10))

        self.tree = ttk.Treeview(list_frame, columns=("Title", "Category", "Images"), show="headings", height=5)
        self.tree.heading("Title", text="Title")
        self.tree.heading("Category", text="Category")
        self.tree.heading("Images", text="Image Count")
        self.tree.column("Title", width=300)
        self.tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        del_btn = ttk.Button(list_frame, text="Delete Selected Entry", command=self.delete_entry)
        del_btn.pack(side=tk.BOTTOM, pady=5, anchor=tk.E)

        ttk.Button(scroll_frame, text="GENERATE WORD DOC REPORT", command=self.generate_report).pack(fill=tk.X, pady=10)

    def _init_help_tab(self):
        help_text_content = """
        REPORT GENERATOR HELP
        =====================
        1. SAVING & LOADING: Use the SESSION MANAGER TAB to save/load your full work.
        2. ADD FINDINGS: Fill out the form. Paste KQL queries in the KQL box.
        3. SCREENSHOTS: Click 'Add Image' to attach evidence.
        4. GENERATE: Click the bottom button to create a .docx file.
        """
        help_text = scrolledtext.ScrolledText(self.help_tab, font=("Courier New", 11))
        help_text.pack(fill=tk.BOTH, expand=True)
        help_text.insert(tk.END, help_text_content)
        help_text.config(state=tk.DISABLED)

    def create_scrollable_text(self, parent, height=5, font=("Segoe UI", 10)):
        frame = ttk.Frame(parent)
        scrollbar = ttk.Scrollbar(frame, orient=tk.VERTICAL)
        text_widget = tk.Text(frame, height=height, font=font, yscrollcommand=scrollbar.set)
        scrollbar.config(command=text_widget.yview)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        text_widget.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        return frame, text_widget

    def add_image_with_caption(self):
        file_path = filedialog.askopenfilename()
        if not file_path: return
        caption = simpledialog.askstring("Image Caption", f"Enter description for:\n{os.path.basename(file_path)}")
        if caption is None: return
        self.current_images.append({'path': file_path, 'caption': caption or "Evidence"})
        self.img_listbox.insert(tk.END, f"{os.path.basename(file_path)} - {caption}")

    def remove_selected_image(self):
        selection = self.img_listbox.curselection()
        if selection:
            idx = selection[0]
            self.img_listbox.delete(idx)
            del self.current_images[idx]

    def add_entry(self):
        title = self.title_entry.get().strip()
        if not title:
            messagebox.showwarning("Validation Error", "Please provide a Challenge Title.")
            return

        entry = {
            "title": title,
            "category": self.cat_combo.get(),
            "description": self.desc_text.get("1.0", tk.END).strip(),
            "methodology": self.method_text.get("1.0", tk.END).strip(),
            "kql_query": self.kql_text.get("1.0", tk.END).strip(),
            "flag": self.flag_entry.get().strip(),
            "images": list(self.current_images)
        }
        self.entries.append(entry)
        self.tree.insert("", tk.END, values=(entry['title'], entry['category'], str(len(entry['images']))))
        self.clear_form()

    def delete_entry(self):
        selected = self.tree.selection()
        if selected:
            item = selected[0]
            idx = self.tree.index(item)
            self.tree.delete(item)
            del self.entries[idx]

    def clear_form(self):
        self.title_entry.delete(0, tk.END)
        self.desc_text.delete("1.0", tk.END)
        self.method_text.delete("1.0", tk.END)
        self.kql_text.delete("1.0", tk.END)
        self.flag_entry.delete(0, tk.END)
        self.current_images = []
        self.img_listbox.delete(0, tk.END)

    def generate_report(self):
        if not self.entries: return
        f = filedialog.asksaveasfilename(defaultextension=".docx")
        if not f: return
        try:
            doc = Document()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("Threat Hunt Report")
            run.bold = True; run.font.size = Pt(24); run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

            p2 = doc.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p2.add_run(f"\nEvent: {self.event_entry.get()}\nAuthor: {self.author_entry.get()}\nDate: {datetime.datetime.now().strftime('%Y-%m-%d')}")
            doc.add_page_break()

            for i, e in enumerate(self.entries, 1):
                doc.add_heading(f"{i}. {e['title']}", level=1)
                doc.add_paragraph(f"Category: {e['category']}")
                doc.add_heading('Description', level=2)
                doc.add_paragraph(e['description'])
                doc.add_heading('Methodology', level=2)
                doc.add_paragraph(e['methodology'])

                if e.get('kql_query'):
                    doc.add_heading('KQL Query', level=3)
                    run = doc.add_paragraph().add_run(e['kql_query'])
                    run.font.name = 'Courier New'

                if e.get('images'):
                    doc.add_heading('Evidence', level=3)
                    for img in e['images']:
                        if os.path.exists(img['path']):
                            doc.add_picture(img['path'], width=Inches(6))
                            doc.add_paragraph(f"Caption: {img['caption']}", style="Caption")

                doc.add_heading('Flag', level=2)
                run = doc.add_paragraph().add_run(e['flag'])
                run.font.color.rgb = RGBColor(255, 0, 0); run.bold = True
                doc.add_page_break()

            doc.save(f)
            messagebox.showinfo("Success", f"Report saved to {f}")
        except Exception as ex:
            messagebox.showerror("Error", str(ex))

    def get_state(self):
        return {
            "entries": self.entries,
            "author": self.author_entry.get(),
            "event": self.event_entry.get()
        }

    def set_state(self, data):
        self.entries = data.get("entries", [])
        self.author_entry.delete(0, tk.END)
        self.author_entry.insert(0, data.get("author", ""))
        self.event_entry.delete(0, tk.END)
        self.event_entry.insert(0, data.get("event", ""))
        for item in self.tree.get_children():
            self.tree.delete(item)
        for e in self.entries:
            self.tree.insert("", tk.END, values=(e['title'], e['category'], str(len(e.get('images', [])))))

# ==========================================
# 4. GUI APPLICATION MAIN
# ==========================================

class UnifiedSOCTool:
    def __init__(self, root):
        self.root = root
        self.root.title("Unified SOC Analyst & Threat Hunter (CTF Edition)")
        self.root.geometry("1400x900")

        self.provider_var = tk.StringVar(value="OpenAI")
        self.api_key_vars = {
            "OpenAI": tk.StringVar(value=os.environ.get("OPENAI_API_KEY", "")),
            "Gemini": tk.StringVar(value=os.environ.get("GEMINI_API_KEY", "")),
            "Claude": tk.StringVar(value=os.environ.get("ANTHROPIC_API_KEY", "")),
        }
        # Backwards-compatible alias — returns the active provider's key
        self.api_key_var = self.api_key_vars["OpenAI"]
        self.workspace_id_var = tk.StringVar(value="")
        self._model_combos = []  # List of (model_var, combo_widget) for provider switching
        self.custom_models = {"OpenAI": [], "Gemini": [], "Claude": []}  # User-added model names

        self.active_flag_var = tk.StringVar(value="General/All") # GLOBAL FLAG FOCUS

        self.th_files = []
        self.th_found_flags = set()
        self.verified_flags_data = [] # List of {title, description, note, focus_id}

        self.th_candidate_queue = []
        self.th_pages_buffer = []
        self.th_current_idx = 0
        self.th_model_var = tk.StringVar(value=DEFAULT_MODEL)

        self.soc_memory = []
        self.soc_last_records = []
        self.soc_logger = None
        self.soc_model_var = tk.StringVar(value=DEFAULT_MODEL)
        self.last_kql = ""
        self.last_error = ""
        self._ai_cancel_event = threading.Event()
        self._ai_running = False
        self._flag_bank_cache = ""

        self.ctf_hints = []
        self.hint_model_var = tk.StringVar(value=DEFAULT_MODEL)
        self.new_hint_flag_var = tk.StringVar(value="Flag 1") # Hint Assignment

        self.ir_file_var = tk.StringVar()
        self.ir_model_var = tk.StringVar(value=DEFAULT_MODEL)
        self.ir_template_default = """
Investigation Template

Report Template
Findings (What did you find)
Investigation Summary (What happened)

Who, What, When, Where, Why, How (Answer as much as possible)
Who - Who was involved?
What - What happened?
When - When did this occur and is it still happening? (Include timezone)
Where - Where in the environment did this happen?
Why - Why did this happen? (If known)
How - How did this happen?

Recommendations: (What steps should be taken to reduce risk or stop the activity?)
"""
        self._setup_ui()

    def _setup_ui(self):
        # --- TOP BAR GLOBAL FOCUS ---
        top_bar = ttk.Frame(self.root, padding=5, relief="raised")
        top_bar.pack(fill="x", side="top")

        ttk.Label(top_bar, text="🚩 Current Investigation Focus:", font=("Segoe UI", 10, "bold")).pack(side="left", padx=5)
        # Scaled to 100 Flags
        flag_opts = ["General/All"] + [f"Flag {i}" for i in range(1, 101)]
        self.flag_selector = ttk.Combobox(top_bar, textvariable=self.active_flag_var, values=flag_opts, state="readonly", width=15)
        self.flag_selector.pack(side="left", padx=5)

        ttk.Label(top_bar, text="(Selects which hints/context the AI uses)").pack(side="left", padx=5)
        # -----------------------------

        self.notebook = ttk.Notebook(self.root)
        self.notebook.pack(fill="both", expand=True, padx=5, pady=5)

        self.tab_config = ttk.Frame(self.notebook)
        self.tab_hints = ttk.Frame(self.notebook)
        self.tab_hunter = ttk.Frame(self.notebook)
        self.tab_soc = ttk.Frame(self.notebook)
        self.tab_flag_bank = ttk.Frame(self.notebook) # NEW FLAG BANK
        self.tab_summary = ttk.Frame(self.notebook)
        self.tab_report = ttk.Frame(self.notebook)
        self.tab_incident = ttk.Frame(self.notebook)
        self.tab_session = ttk.Frame(self.notebook)
        self.tab_guide = ttk.Frame(self.notebook)

        self.notebook.add(self.tab_config, text="⚙️ Configuration")
        self.notebook.add(self.tab_hints, text="🧩 Flag Hints")
        self.notebook.add(self.tab_hunter, text="🕵️ Threat Hunter")
        self.notebook.add(self.tab_soc, text="🛡️ Azure SOC Agent")
        self.notebook.add(self.tab_flag_bank, text="🏦 Flag Bank (Context)") # NEW
        self.notebook.add(self.tab_summary, text="🏆 Flag Summary")
        self.notebook.add(self.tab_report, text="📝 Report Editor")
        self.notebook.add(self.tab_incident, text="📑 Incident Report Generator")
        self.notebook.add(self.tab_session, text="💾 Session Manager")
        self.notebook.add(self.tab_guide, text="📖 How-To Guide")

        self._setup_config_tab()
        self._setup_hints_tab()
        self._setup_hunter_tab()
        self._setup_soc_tab()
        self._setup_flag_bank_tab() # NEW
        self._setup_summary_tab()
        self._setup_incident_tab()
        self._setup_session_tab()

        self._setup_guide_tab()
        self.reporter = ThreatHuntReporterTab(self.tab_report)

    # -------------------------------------------------------------------------
    # CONFIGURATION
    # -------------------------------------------------------------------------
    def _setup_config_tab(self):
        # AI Provider Selection
        provider_frame = ttk.LabelFrame(self.tab_config, text="AI Provider", padding=20)
        provider_frame.pack(fill="x", padx=20, pady=10)

        ttk.Label(provider_frame, text="Active Provider:").grid(row=0, column=0, sticky="w", pady=10)
        provider_combo = ttk.Combobox(provider_frame, textvariable=self.provider_var, values=PROVIDER_OPTIONS, state="readonly", width=20)
        provider_combo.grid(row=0, column=1, padx=10, sticky="w")
        provider_combo.bind("<<ComboboxSelected>>", lambda e: self._on_provider_changed())

        # Custom Model Entry
        model_frame = ttk.LabelFrame(self.tab_config, text="Custom AI Models", padding=20)
        model_frame.pack(fill="x", padx=20, pady=10)

        ttk.Label(model_frame, text="Add a model ID not yet in the built-in list (e.g. gpt-6, claude-opus-5):").grid(row=0, column=0, columnspan=3, sticky="w", pady=(0, 5))

        ttk.Label(model_frame, text="Model ID:").grid(row=1, column=0, sticky="w", pady=5)
        self.custom_model_entry = ttk.Entry(model_frame, width=40)
        self.custom_model_entry.grid(row=1, column=1, padx=10, sticky="w")
        ttk.Button(model_frame, text="Add to Current Provider", command=self._add_custom_model).grid(row=1, column=2, padx=10)

        ttk.Label(model_frame, text="Custom models for current provider:").grid(row=2, column=0, sticky="w", pady=(10, 0))
        self.custom_models_label = ttk.Label(model_frame, text="(none)", foreground="gray")
        self.custom_models_label.grid(row=2, column=1, columnspan=2, sticky="w", padx=10, pady=(10, 0))

        ttk.Button(model_frame, text="Remove Selected Custom Model", command=self._remove_custom_model).grid(row=3, column=2, padx=10, pady=5)
        self.remove_model_entry = ttk.Entry(model_frame, width=40)
        self.remove_model_entry.grid(row=3, column=1, padx=10, sticky="w")
        ttk.Label(model_frame, text="Model to remove:").grid(row=3, column=0, sticky="w", pady=5)

        # API Keys for all providers
        keys_frame = ttk.LabelFrame(self.tab_config, text="API Keys", padding=20)
        keys_frame.pack(fill="x", padx=20, pady=10)

        ttk.Label(keys_frame, text="OpenAI API Key:").grid(row=0, column=0, sticky="w", pady=5)
        ttk.Entry(keys_frame, textvariable=self.api_key_vars["OpenAI"], width=60, show="*").grid(row=0, column=1, padx=10)

        ttk.Label(keys_frame, text="Gemini API Key:").grid(row=1, column=0, sticky="w", pady=5)
        ttk.Entry(keys_frame, textvariable=self.api_key_vars["Gemini"], width=60, show="*").grid(row=1, column=1, padx=10)
        ttk.Label(keys_frame, text="(pip install google-generativeai)").grid(row=1, column=2, sticky="w")

        ttk.Label(keys_frame, text="Claude API Key:").grid(row=2, column=0, sticky="w", pady=5)
        ttk.Entry(keys_frame, textvariable=self.api_key_vars["Claude"], width=60, show="*").grid(row=2, column=1, padx=10)
        ttk.Label(keys_frame, text="(pip install anthropic)").grid(row=2, column=2, sticky="w")

        # Azure settings
        frame = ttk.LabelFrame(self.tab_config, text="Azure Settings", padding=20)
        frame.pack(fill="x", padx=20, pady=10)

        ttk.Label(frame, text="Log Analytics Workspace ID:").grid(row=0, column=0, sticky="w", pady=10)
        ttk.Entry(frame, textvariable=self.workspace_id_var, width=60).grid(row=0, column=1, padx=10)

        ttk.Label(frame, text="(Required for SOC Agent tab only)").grid(row=0, column=2, sticky="w")
        ttk.Button(frame, text="Save / Validate", command=lambda: messagebox.showinfo("Info", "Settings ready in memory.")).grid(row=1, column=1, pady=20)

    def _get_provider(self):
        """Return the currently selected AI provider name."""
        return self.provider_var.get()

    def _get_api_key(self):
        """Return the API key for the currently selected provider."""
        return self.api_key_vars[self.provider_var.get()].get()

    def _get_models_for_provider(self, provider=None):
        """Return built-in + custom models for the given provider."""
        if provider is None:
            provider = self.provider_var.get()
        built_in = list(PROVIDER_MODELS.get(provider, MODEL_OPTIONS))
        custom = self.custom_models.get(provider, [])
        return built_in + custom

    def _on_provider_changed(self):
        """Called when the user changes the AI provider dropdown."""
        provider = self.provider_var.get()
        default_model = PROVIDER_DEFAULTS.get(provider, "gpt-4o")
        models = self._get_models_for_provider(provider)

        # Update all model dropdowns
        for model_var, combo in self._model_combos:
            model_var.set(default_model)
            combo['values'] = models

        self._update_custom_models_label()

    def _add_custom_model(self):
        """Add a user-entered model ID to the current provider's model list."""
        model_id = self.custom_model_entry.get().strip()
        if not model_id:
            messagebox.showwarning("No Model ID", "Please enter a model ID.")
            return

        provider = self.provider_var.get()
        all_models = self._get_models_for_provider(provider)

        if model_id in all_models:
            messagebox.showinfo("Already Exists", f"'{model_id}' is already available for {provider}.")
            return

        self.custom_models[provider].append(model_id)
        updated_models = self._get_models_for_provider(provider)

        # Update all model dropdowns to include the new model
        for model_var, combo in self._model_combos:
            combo['values'] = updated_models

        self.custom_model_entry.delete(0, "end")
        self._update_custom_models_label()
        messagebox.showinfo("Model Added", f"'{model_id}' added to {provider}. You can now select it in any tab's model dropdown.")

    def _remove_custom_model(self):
        """Remove a custom model from the current provider's list."""
        model_id = self.remove_model_entry.get().strip()
        if not model_id:
            messagebox.showwarning("No Model ID", "Please enter the model ID to remove.")
            return

        provider = self.provider_var.get()
        if model_id not in self.custom_models.get(provider, []):
            messagebox.showwarning("Not Found", f"'{model_id}' is not a custom model for {provider}.")
            return

        self.custom_models[provider].remove(model_id)
        updated_models = self._get_models_for_provider(provider)

        for model_var, combo in self._model_combos:
            combo['values'] = updated_models
            # If the removed model was selected, reset to default
            if model_var.get() == model_id:
                model_var.set(PROVIDER_DEFAULTS.get(provider, "gpt-4o"))

        self.remove_model_entry.delete(0, "end")
        self._update_custom_models_label()
        messagebox.showinfo("Model Removed", f"'{model_id}' removed from {provider}.")

    def _update_custom_models_label(self):
        """Update the label showing custom models for the current provider."""
        provider = self.provider_var.get()
        custom = self.custom_models.get(provider, [])
        if custom:
            self.custom_models_label.config(text=", ".join(custom), foreground="black")
        else:
            self.custom_models_label.config(text="(none)", foreground="gray")

    # -------------------------------------------------------------------------
    # HINTS (UPDATED)
    # -------------------------------------------------------------------------
    def _setup_hints_tab(self):
        top_frame = ttk.LabelFrame(self.tab_hints, text="Add New Hint", padding=10)
        top_frame.pack(fill="x", padx=10, pady=5)

        # New Flag ID Selection (Scaled to 100)
        row1 = ttk.Frame(top_frame)
        row1.grid(row=0, column=0, columnspan=3, sticky="ew", pady=5)
        ttk.Label(row1, text="Assign to Flag:").pack(side="left")
        ttk.Combobox(row1, textvariable=self.new_hint_flag_var, values=[f"Flag {i}" for i in range(1, 101)], width=10).pack(side="left", padx=5)

        ttk.Label(top_frame, text="Hint Text (Multiline):").grid(row=1, column=0, sticky="nw")
        self.hint_input = scrolledtext.ScrolledText(top_frame, width=80, height=4, font=("Segoe UI", 10))
        self.hint_input.grid(row=1, column=1, padx=5, pady=5)

        opts_frame = ttk.Frame(top_frame)
        opts_frame.grid(row=1, column=2, sticky="nw", padx=5)

        ttk.Label(opts_frame, text="Model:").pack(anchor="w")
        hint_model_combo = ttk.Combobox(opts_frame, textvariable=self.hint_model_var, values=self._get_models_for_provider(), state="readonly", width=25)
        hint_model_combo.pack(fill="x")
        self._model_combos.append((self.hint_model_var, hint_model_combo))
        ttk.Button(opts_frame, text="Analyze Hint", command=self.hint_analyze).pack(fill="x", pady=5)

        bot_frame = ttk.LabelFrame(self.tab_hints, text="Active Hints & Generated Context", padding=10)
        bot_frame.pack(fill="both", expand=True, padx=10, pady=5)

        self.hint_display = scrolledtext.ScrolledText(bot_frame, font=("Consolas", 10), state="disabled")
        self.hint_display.pack(fill="both", expand=True)

    def hint_analyze(self):
        hint_text = self.hint_input.get("1.0", "end-1c").strip()
        flag_id = self.new_hint_flag_var.get()
        if not hint_text: return
        self.hint_display.config(state="normal")
        self.hint_display.insert("end", f"\n[*] Analyzing Hint for {flag_id}: '{hint_text[:50]}...'...\n")
        self.hint_display.config(state="disabled")
        threading.Thread(target=self._hint_thread, args=(hint_text, flag_id), daemon=True).start()

    def _hint_thread(self, hint_text, flag_id):
        try:
            model = self.hint_model_var.get()
            prompt = f"""
            You are a CTF Assistant. The user has a hint for {flag_id}: "{hint_text}".
            1. Explain what artifacts or logs to look for based on this hint.
            2. Generate a specific KQL query to find it in Azure Sentinel.
            Return JSON: {{ "clue": "...", "kql": "..." }}
            """
            content = ai_chat_completion(
                self._get_provider(), self._get_api_key(), model,
                [{"role": "user", "content": prompt}], json_mode=True
            )
            data = json.loads(content)
            entry = {"id": flag_id, "hint": hint_text, "clue": data.get("clue"), "kql": data.get("kql")}
            self.ctf_hints.append(entry)
            self.root.after(0, lambda: self._update_hint_ui(entry))
        except Exception as e:
            print(f"Hint Error: {e}")

    def _update_hint_ui(self, entry):
        msg = f"\n✅ HINT ADDED ({entry['id']}):\nHint: {entry['hint']}\nAI Clue: {entry['clue']}\nSuggested KQL: {entry['kql']}\n" + "-"*50 + "\n"
        self.hint_display.config(state="normal")
        self.hint_display.insert("end", msg)
        self.hint_display.see("end")
        self.hint_display.config(state="disabled")
        self.hint_input.delete("1.0", "end")
        messagebox.showinfo("Success", "Hint added to AI Memory.")

    # -------------------------------------------------------------------------
    # HELPER: GET ACTIVE HINTS & CONTEXT
    # -------------------------------------------------------------------------
    def get_active_hints(self):
        focus = self.active_flag_var.get()
        if focus == "General/All":
            return self.ctf_hints
        return [h for h in self.ctf_hints if h.get("id") == focus]

    def _get_incident_context_for_kql(self):
        """Generates context from verified flags AND FLAG BANK for AI."""
        ctx = ""
        # 1. Verified Flags
        if self.verified_flags_data:
            ctx += "CONFIRMED FINDINGS/FLAGS (Use these to correlate/pivot):\n"
            for f in self.verified_flags_data:
                ctx += f"- Found: {f.get('title', 'Unknown')} | Note: {f.get('note', '')}\n"

        # 2. Flag Bank (Current Understanding) - use cached copy for thread safety
        bank_text = self._flag_bank_cache
        if bank_text:
            ctx += f"\nCURRENT INCIDENT UNDERSTANDING (Narrative):\n{bank_text}\n"

        return ctx

    # -------------------------------------------------------------------------
    # FLAG BANK (NEW)
    # -------------------------------------------------------------------------
    def _setup_flag_bank_tab(self):
        frame = ttk.LabelFrame(self.tab_flag_bank, text="Current Incident Understanding (AI Generated)", padding=10)
        frame.pack(fill="both", expand=True, padx=10, pady=10)

        self.flag_bank_text = scrolledtext.ScrolledText(frame, font=("Segoe UI", 11), state="normal")
        self.flag_bank_text.pack(fill="both", expand=True)

        btn_frame = ttk.Frame(frame)
        btn_frame.pack(fill="x", pady=5)

        ttk.Button(btn_frame, text="🔄 Update Understanding with AI", command=self.update_flag_bank_ai).pack(side="right")
        ttk.Button(btn_frame, text="📤 Export Narrative to Incident Report", command=self.export_flag_bank_to_ir).pack(side="right", padx=10)

    def update_flag_bank_ai(self):
        if not self.verified_flags_data and not self.ctf_hints:
            return
        if not self._get_api_key():
            return

        # Get data on main thread
        flags_data = list(self.verified_flags_data)
        hints_data = list(self.ctf_hints)

        self.flag_bank_text.insert("end", "\n[System] Updating narrative...\n")
        self.flag_bank_text.see("end")

        threading.Thread(target=self._flag_bank_ai_thread, args=(flags_data, hints_data), daemon=True).start()

    def _flag_bank_ai_thread(self, flags, hints):
        try:
            context_str = "VERIFIED FACTS:\n" + "\n".join([f"- {f['title']}: {f['note']}" for f in flags])
            context_str += "\n\nKNOWN HINTS:\n" + "\n".join([f"- {h['hint']}" for h in hints])

            prompt = f"""
            You are a Lead Incident Responder.
            Based on the following Verified Facts and Known Hints, write a cohesive "Current Understanding of the Incident" narrative.

            - Connect the dots between flags.
            - Highlight what is confirmed vs what is suspected (hints).
            - Keep it professional and chronological if possible.

            DATA:
            {context_str}
            """

            provider = self._get_provider()
            model = PROVIDER_DEFAULTS.get(provider, DEFAULT_MODEL)
            narrative = ai_chat_completion(
                provider, self._get_api_key(), model,
                [{"role": "user", "content": prompt}]
            )

            self.root.after(0, lambda: self._update_flag_bank_ui(narrative))

        except Exception as e:
            print(f"Bank AI Error: {e}")

    def _update_flag_bank_ui(self, text):
        self.flag_bank_text.delete("1.0", "end")
        self.flag_bank_text.insert("1.0", text)
        self._flag_bank_cache = text

    def export_flag_bank_to_ir(self):
        """Export Flag Bank content to Incident Report template."""
        narrative = self.flag_bank_text.get("1.0", "end").strip()
        if not narrative:
            messagebox.showwarning("Empty", "Flag Bank is empty.")
            return

        # Switch to IR tab
        self.notebook.select(self.tab_incident)

        # Append to template
        current_template = self.ir_template_text.get("1.0", "end")
        if "CONTEXT FROM FLAG BANK" not in current_template:
            self.ir_template_text.insert("end", f"\n\n=== CONTEXT FROM FLAG BANK ===\n{narrative}\n")
            messagebox.showinfo("Exported", "Flag Bank narrative appended to Report Template.")
        else:
            # Maybe replace? For now just append or warn.
            messagebox.showinfo("Info", "Flag Bank context appears to be already present in the template.")

    # -------------------------------------------------------------------------
    # THREAT HUNTER
    # -------------------------------------------------------------------------
    def _setup_hunter_tab(self):
        top_frame = ttk.LabelFrame(self.tab_hunter, text="Input Data & Settings", padding=10)
        top_frame.pack(fill="x", padx=10, pady=5)

        self.file_listbox = tk.Listbox(top_frame, height=4, width=60)
        self.file_listbox.pack(side="left", fill="x", expand=True, padx=5)

        opts_frame = ttk.Frame(top_frame)
        opts_frame.pack(side="right", fill="y")

        ttk.Label(opts_frame, text="AI Model:").pack(anchor="w")
        th_model_combo = ttk.Combobox(opts_frame, textvariable=self.th_model_var, values=self._get_models_for_provider(), state="readonly", width=25)
        th_model_combo.pack(fill="x", pady=2)
        self._model_combos.append((self.th_model_var, th_model_combo))

        ttk.Button(opts_frame, text="Add Files...", command=self.th_add_files).pack(fill="x", pady=2)
        ttk.Button(opts_frame, text="Clear List", command=self.th_clear_files).pack(fill="x", pady=2)
        self.th_start_btn = ttk.Button(opts_frame, text="▶ START HUNT", command=self.th_start_hunt)
        self.th_start_btn.pack(fill="x", pady=(10, 2))
        self.th_stop_btn = ttk.Button(opts_frame, text="■ STOP HUNT", command=self.th_stop_hunt, state="disabled")
        self.th_stop_btn.pack(fill="x", pady=2)

        mid_frame = ttk.LabelFrame(self.tab_hunter, text="Potential Flag / Finding", padding=10)
        mid_frame.pack(fill="both", expand=True, padx=10, pady=5)

        self.th_editor = scrolledtext.ScrolledText(mid_frame, height=15, font=("Consolas", 10))
        self.th_editor.pack(fill="both", expand=True)

        action_frame = ttk.Frame(mid_frame)
        action_frame.pack(fill="x", pady=5)
        ttk.Button(action_frame, text="✅ Verify & Save", command=self.th_verify).pack(side="left", padx=5)
        ttk.Button(action_frame, text="❌ Discard", command=self.th_discard).pack(side="left", padx=5)

        # New Manual Flag Button
        ttk.Button(action_frame, text="➕ Manually Add Flag", command=self.th_manual_add).pack(side="left", padx=20)

        self.th_status_lbl = ttk.Label(action_frame, text="Status: Idle", foreground="blue")
        self.th_status_lbl.pack(side="right", padx=10)

        bot_frame = ttk.LabelFrame(self.tab_hunter, text="Verified Flags History", padding=10)
        bot_frame.pack(fill="x", padx=10, pady=5)
        self.th_history = scrolledtext.ScrolledText(bot_frame, height=6, state="disabled")
        self.th_history.pack(fill="both", expand=True)

    def th_add_files(self):
        files = filedialog.askopenfilenames(filetypes=[("All Supported", "*.pdf *.txt *.log *.jsonl"), ("PDF", "*.pdf"), ("Text", "*.txt")])
        for f in files:
            if f not in self.th_files:
                self.th_files.append(f)
                self.file_listbox.insert("end", os.path.basename(f))

    def th_clear_files(self):
        self.th_files = []
        self.file_listbox.delete(0, "end")

    def th_log_history(self, msg):
        self.th_history.config(state="normal")
        self.th_history.insert("end", msg + "\n")
        self.th_history.see("end")
        self.th_history.config(state="disabled")

    def th_start_hunt(self):
        if not self.th_files:
            messagebox.showerror("Error", "No files loaded.")
            return
        if not self._get_api_key():
            messagebox.showerror("Error", "Missing API Key in Configuration tab.")
            return

        self.th_pages_buffer = []
        for f in self.th_files:
            self.th_status_lbl.config(text=f"Reading {os.path.basename(f)}...")
            self.root.update()
            pages = extract_text_from_file(f)
            self.th_pages_buffer.extend(pages)

        if not self.th_pages_buffer:
            messagebox.showerror("Error", "No text could be extracted.")
            return

        self.th_current_idx = 0
        self.th_candidate_queue = []
        self.th_status_lbl.config(text=f"Processing {len(self.th_pages_buffer)} pages/chunks...")
        self._ai_cancel_event.clear()
        self._set_ai_running(True)
        threading.Thread(target=self.th_process_batch_thread, daemon=True).start()

    def th_process_batch_thread(self):
        while self.th_current_idx < len(self.th_pages_buffer):
            if self._ai_cancelled():
                self.root.after(0, lambda: self.th_status_lbl.config(text="Hunt stopped by user."))
                self._set_ai_running(False)
                return

            if self.th_candidate_queue:
                self.root.after(0, self.th_show_candidate)
                return

            end_idx = min(self.th_current_idx + 3, len(self.th_pages_buffer))
            batch_text = "\n".join(self.th_pages_buffer[self.th_current_idx:end_idx])
            self.th_current_idx = end_idx

            try:
                model = self.th_model_var.get()

                # USE ACTIVE HINTS
                active_hints = self.get_active_hints()
                current_focus = self.active_flag_var.get()

                hints_ctx = ""
                if active_hints:
                    hints_ctx = f"CTF HINTS (FOCUS: {current_focus}):\n" + "\n".join([f"[{h.get('id','?')}] {h['hint']}" for h in active_hints])

                prompt = f"""You are a CTF Flag Hunter. Analyze this log excerpt.
Your goal is to find the EXACT FLAG ANSWER - the specific value, string, artifact, IP, username, hash, or flag{{...}} that answers the CTF challenge or investigation question.

{hints_ctx}
Already found (ignore these): {json.dumps(list(self.th_found_flags))}

IMPORTANT: Do NOT just describe threats generically. Extract the EXACT answer value.
Return JSON: {{ "findings": [ {{ "title": "short name", "description": "why this is the answer", "flag_answer": "THE EXACT VALUE e.g. flag{{abc123}} or 192.168.1.5 or malware.exe", "severity": "High", "evidence": "raw log line proving it" }} ] }}
If no flags found, return: {{ "findings": [] }}

LOGS:
{batch_text[:15000]}"""
                content = ai_chat_completion(
                    self._get_provider(), self._get_api_key(), model,
                    [{"role": "user", "content": prompt}], json_mode=True
                )
                data = json.loads(content)
                findings = data.get("findings", [])
                for f in findings:
                    if f.get("title") not in self.th_found_flags:
                        self.th_candidate_queue.append(f)
            except Exception as e:
                print(f"AI Error: {e}")
        self._set_ai_running(False)
        self.root.after(0, lambda: messagebox.showinfo("Done", "Hunt Complete."))

    def th_show_candidate(self):
        if not self.th_candidate_queue:
            self._ai_cancel_event.clear()
            self._set_ai_running(True)
            threading.Thread(target=self.th_process_batch_thread, daemon=True).start()
            return
        finding = self.th_candidate_queue.pop(0)
        formatted = AzureSentinelFormatter.format_log(finding)
        self.th_editor.delete("1.0", "end")
        self.th_editor.insert("1.0", formatted)
        self.th_status_lbl.config(text="Waiting for Verification...")

    def th_manual_add(self):
        """Allows manual entry of a flag."""
        title = simpledialog.askstring("Manual Flag", "Enter Flag Title/Threat Name:")
        if not title: return

        description = simpledialog.askstring("Manual Flag", "Enter brief description (optional):")

        finding = {
            "title": title,
            "description": description if description else "Manually added finding.",
            "timestamp": datetime.datetime.now().isoformat(),
            "evidence": "Manual Entry"
        }

        # Populate editor and let verify logic handle the rest
        formatted = AzureSentinelFormatter.format_log(finding)
        self.th_editor.delete("1.0", "end")
        self.th_editor.insert("1.0", formatted)
        self.th_status_lbl.config(text="Waiting for Verification of Manual Entry...")

    def th_verify(self):
        try:
            content = self.th_editor.get("1.0", "end").strip()
            if not content: return
            data = json.loads(content)
            title = data.get("AlertName", data.get("title", "Unknown"))
            description = data.get("Description", "")
            flag_answer = data.get("FlagAnswer", data.get("flag_answer", ""))

            # --- NEW AUTO-NOTE LOGIC ---
            self.th_status_lbl.config(text="Drafting Flag Note (AI)...")
            self.root.update()

            focus_id = self.active_flag_var.get()

            # Find relevant hint
            relevant_hint = "No specific hint linked."
            if focus_id != "General/All":
                for h in self.ctf_hints:
                    if h.get('id') == focus_id:
                        relevant_hint = h.get('hint', '')
                        break

            # Draft Note
            suggested_note = ""
            try:
                if not self._get_api_key():
                    raise ValueError("API key not configured")
                # Use the findings specific evidence if available
                evidence = data.get("Evidence", str(data))

                prompt = f"""Context: CTF Investigation. Focus: {focus_id}.
Hint provided: "{relevant_hint}"

The analyst found this:
Title: {title}
Flag Answer: {flag_answer}
Description: {description}
Evidence/Data: {evidence}

Task: Write a very brief (1-sentence) note stating the EXACT flag answer or artifact value.
If a flag_answer is provided, include it verbatim in your note.
Example: "The flag is flag{{abc123}}" or "The malicious IP is 192.168.1.5"
"""

                provider = self._get_provider()
                # Use a fast model for note generation
                fast_models = {"OpenAI": "gpt-4o-mini", "Gemini": "gemini-2.0-flash", "Claude": "claude-haiku-4-5-20251001"}
                fast_model = fast_models.get(provider, PROVIDER_DEFAULTS.get(provider, DEFAULT_MODEL))
                suggested_note = ai_chat_completion(
                    provider, self._get_api_key(), fast_model,
                    [{"role": "user", "content": prompt}], max_tokens=60
                ).strip()
            except Exception as ai_e:
                print(f"Note Gen Error: {ai_e}")
                if flag_answer:
                    suggested_note = f"Flag answer: {flag_answer}"
                else:
                    suggested_note = f"Found {title}. Evidence: {data.get('Evidence', 'See logs')}"

            # Dialog
            note = simpledialog.askstring(
                "Flag Answer/Note",
                f"Verified '{title}'.\nEdit the generated note below:",
                initialvalue=suggested_note
            )

            self.th_status_lbl.config(text="Idle")
            if note is None: return # Cancelled
            if not note: note = "No specific note provided."
            # ---------------------------

            # Save
            self.th_found_flags.add(title)
            self.th_log_history(f"[FLAG] {title}")

            # Store in Summary Data
            self.verified_flags_data.append({
                "title": title,
                "description": description,
                "note": note,
                "focus_id": focus_id
            })
            self._update_summary_display()

            # --- AUTO TRIGGER INCIDENT REPORT & FLAG BANK ---
            self.root.after(500, self.auto_generate_incident_report)
            self.root.after(1000, self.update_flag_bank_ai) # Trigger Bank update
            # ------------------------------------

            self.th_editor.delete("1.0", "end")
            self.th_show_candidate()
        except Exception as e:
            messagebox.showerror("Error", f"Invalid JSON or Error: {e}")

    def th_discard(self):
        self.th_editor.delete("1.0", "end")
        self.th_show_candidate()

    # -------------------------------------------------------------------------
    # SOC AGENT
    # -------------------------------------------------------------------------
    def _setup_soc_tab(self):
        console_frame = ttk.LabelFrame(self.tab_soc, text="Agent Console", padding=5)
        console_frame.pack(fill="both", expand=True, padx=10, pady=5)
        self.soc_console = scrolledtext.ScrolledText(console_frame, state="disabled", height=25, font=("Consolas", 9), bg="black", fg="lime")
        self.soc_console.pack(fill="both", expand=True)
        self.logger = SessionLogger(self.soc_console)

        ctrl_frame = ttk.Frame(self.tab_soc)
        ctrl_frame.pack(fill="x", padx=10, pady=10)

        top_ctrl = ttk.Frame(ctrl_frame)
        top_ctrl.pack(fill="x", pady=2)
        ttk.Label(top_ctrl, text="Instruction / Query:").pack(side="left")
        ttk.Label(top_ctrl, text="Model:").pack(side="right")
        soc_model_combo = ttk.Combobox(top_ctrl, textvariable=self.soc_model_var, values=self._get_models_for_provider(), state="readonly", width=25)
        soc_model_combo.pack(side="right")
        self._model_combos.append((self.soc_model_var, soc_model_combo))

        self.soc_prompt_text = tk.Text(ctrl_frame, height=3, font=("Consolas", 10), wrap="word")
        self.soc_prompt_text.pack(fill="x", pady=2)

        btn_box = ttk.Frame(ctrl_frame)
        btn_box.pack(fill="x", pady=5)
        self.soc_run_btn = ttk.Button(btn_box, text="Run AI Investigation", command=self.soc_run_ai)
        self.soc_run_btn.pack(side="left", padx=2)
        self.soc_stop_btn = ttk.Button(btn_box, text="Stop AI", command=self.soc_stop_ai, state="disabled")
        self.soc_stop_btn.pack(side="left", padx=2)
        ttk.Button(btn_box, text="Generate KQL Only", command=self.soc_gen_kql).pack(side="left", padx=2)
        ttk.Button(btn_box, text="Open Manual KQL Editor", command=self.soc_open_manual_kql).pack(side="left", padx=2)
        ttk.Button(btn_box, text="💾 Export Log to File", command=self.soc_export_log).pack(side="right", padx=2)
        ttk.Button(btn_box, text="🩹 Self-Heal Last KQL", command=self.soc_self_heal).pack(side="right", padx=2)

    def _soc_precheck(self):
        if not self._get_api_key() or not self.workspace_id_var.get():
            messagebox.showerror("Config Error", "Please set API Key and Workspace ID in Configuration Tab.")
            return False
        return True

    def _set_ai_running(self, running):
        """Toggle AI running state and update button availability."""
        self._ai_running = running
        def _update():
            if running:
                self.soc_run_btn.config(state="disabled")
                self.soc_stop_btn.config(state="normal")
                self.th_start_btn.config(state="disabled")
                self.th_stop_btn.config(state="normal")
            else:
                self.soc_run_btn.config(state="normal")
                self.soc_stop_btn.config(state="disabled")
                self.th_start_btn.config(state="normal")
                self.th_stop_btn.config(state="disabled")
        self.root.after(0, _update)

    def soc_stop_ai(self):
        """Signal all running AI threads to stop."""
        if self._ai_running:
            self._ai_cancel_event.set()
            self.soc_print("Stopping AI processing... (will halt after current API call completes)")

    def th_stop_hunt(self):
        """Stop the Threat Hunter AI processing."""
        if self._ai_running:
            self._ai_cancel_event.set()
            self.th_status_lbl.config(text="Status: Stopping...")
            self.th_log_history("[STOPPED] Hunt stopped by user.")

    def _ai_cancelled(self):
        """Check if cancellation has been requested."""
        return self._ai_cancel_event.is_set()

    def soc_print(self, msg):
        # FIX: Ensure thread safety by using root.after
        self.root.after(0, lambda: self.logger.write(msg + "\n"))

    def _process_results_in_batches(self, results, table_name, model):
        """Unified logic to process records in batches with detailed feedback."""
        if not results:
            self.soc_print("No records to process.")
            return

        # --- AUTO-SAVE LOGIC ---
        try:
            ts = int(time.time())
            filename = f"azure_raw_logs_{ts}.jsonl"

            # FIX: Write to file in thread, but update GUI safely
            with open(filename, "w", encoding="utf-8") as f:
                for r in results:
                    f.write(json.dumps(r, default=str) + "\n")

            self.soc_print(f"{Fore.CYAN}Saved {len(results)} raw logs to {filename} and added to Threat Hunter.{Fore.RESET}")

            # FIX: Update file listbox safely
            if filename not in self.th_files:
                def _add_file():
                    self.th_files.append(filename)
                    self.file_listbox.insert("end", os.path.basename(filename))
                self.root.after(0, _add_file)

        except Exception as e:
            self.soc_print(f"{Fore.RED}Auto-save error: {e}{Fore.RESET}")

        # --- BATCH PROCESSING ---
        total_records = len(results)
        self.soc_print(f"{total_records} records found.")

        batch_size = 50
        total_batches = math.ceil(total_records / batch_size)

        self.soc_print(f"Processing in batches of {batch_size}...")

        provider = self._get_provider()
        api_key = self._get_api_key()
        all_findings = []

        # Use Active Focus Hints
        active_hints = self.get_active_hints()
        focus = self.active_flag_var.get()

        for i in range(total_batches):
            if self._ai_cancelled():
                self.soc_print(f"{Fore.YELLOW}--- AI PROCESSING STOPPED (completed {i}/{total_batches} batches) ---{Fore.RESET}")
                break

            start_idx = i * batch_size
            end_idx = min((i + 1) * batch_size, total_records)
            batch_records = results[start_idx:end_idx]

            self.soc_print(f"Batch {i+1} of {total_batches}...")

            current_ignore_list = set(self.th_found_flags)
            for f in all_findings:
                current_ignore_list.add(f.get('title'))

            try:
                batch_findings = hunt_on_records(provider, api_key, batch_records, table_name, model, active_hints, current_ignore_list, focus)

                new_count = 0
                if batch_findings:
                    for f in batch_findings:
                        if f.get('title') not in current_ignore_list:
                            all_findings.append(f)
                            new_count += 1

                self.soc_print(f"  + {new_count} flag(s)/answer(s) found.")

            except Exception as e:
                self.soc_print(f"{Fore.RED}  [!] Error processing batch: {e}{Fore.RESET}")

        self.soc_print("End of Batches.")

        if all_findings:
            self.soc_print(f"{Fore.RED}!!! FLAGS / ANSWERS FOUND ({len(all_findings)}) !!!{Fore.RESET}")
            for f in all_findings:
                answer = f.get('flag_answer', '')
                title = f.get('title', 'Unknown')
                if answer:
                    self.soc_print(f"  FLAG: {title} -> {Fore.CYAN}{answer}{Fore.RESET}")
                else:
                    self.soc_print(f"  - {title}: {f.get('description', '')[:100]}")

            # Auto-push findings to Threat Hunter candidate queue
            def _push_to_hunter():
                for f in all_findings:
                    if f.get('title') not in self.th_found_flags:
                        self.th_candidate_queue.append(f)
                self.th_show_candidate()
            self.root.after(0, _push_to_hunter)
            self.soc_print("Findings sent to Threat Hunter tab for verification.")
        else:
            self.soc_print("No flags or answers found in logs.")

    def soc_run_ai(self):
        if not self._soc_precheck(): return
        if self._ai_running:
            messagebox.showinfo("Busy", "AI is already running. Stop it first before starting a new investigation.")
            return
        user_input = self.soc_prompt_text.get("1.0", "end").strip()
        if not user_input:
            messagebox.showinfo("Input", "Please enter instructions above.")
            return
        self._ai_cancel_event.clear()
        self._set_ai_running(True)
        threading.Thread(target=self._soc_ai_thread, args=(user_input,), daemon=True).start()

    def _soc_ai_thread(self, user_input):
        self.soc_print(f"{Fore.CYAN}--- STARTING AI INVESTIGATION ---{Fore.RESET}")
        self.soc_print(f"Goal: {user_input}")

        try:
            provider = self._get_provider()
            api_key = self._get_api_key()
            law_client = LogsQueryClient(credential=DefaultAzureCredential())
            model = self.soc_model_var.get()

            # Use active hints
            active_hints = self.get_active_hints()
            focus = self.active_flag_var.get()

            # Get Accumulated Report Context
            incident_ctx = self._get_incident_context_for_kql()

            if self._ai_cancelled():
                self.soc_print(f"{Fore.YELLOW}--- AI INVESTIGATION STOPPED ---{Fore.RESET}")
                return

            # --- SMART MODE vs SAFE MODE LOGIC ---
            # 1. Ask AI for query
            ctx = get_query_context(provider, api_key, user_input, model, self.soc_memory, active_hints, focus, incident_ctx)

            if self._ai_cancelled():
                self.soc_print(f"{Fore.YELLOW}--- AI INVESTIGATION STOPPED ---{Fore.RESET}")
                return

            # 2. Check if query is valid
            kql = ctx.get('kql_query', '')
            table_name = ctx.get('table_name', 'Unknown')

            if not kql or len(kql) < 10:
                self.soc_print(f"{Fore.YELLOW}AI failed to generate a complex query. Using fallback.{Fore.RESET}")
                # Fallback: Simple search
                clean_input = user_input.replace('"', '').replace("'", "")
                kql = f"search \"{clean_input}\" | take 100"
                table_name = "Search"
            else:
                self.soc_print(f"{Fore.YELLOW}AI Generated KQL:{Fore.RESET}")

            self.soc_print(f"{kql}")

            # 3. Execute
            self.soc_print("Executing query...")

            # Default to 1 year (8760h) for AI queries
            time_range = ctx.get('time_range_hours', 8760)

            results = execute_kql(law_client, self.workspace_id_var.get(), kql, time_range)

            if self._ai_cancelled():
                self.soc_print(f"{Fore.YELLOW}--- AI INVESTIGATION STOPPED ---{Fore.RESET}")
                return

            if results:
                self.soc_memory.append({"user_input": user_input, "kql_query": kql})
                self.soc_last_records = results
                self._process_results_in_batches(results, table_name, model)
            else:
                self.last_kql = kql
                self.last_error = "0 records found."
                self.soc_print(f"{Fore.GREEN}0 records found.{Fore.RESET}")

        except Exception as e:
            self.soc_print(f"{Fore.RED}Error: {e}{Fore.RESET}")
        finally:
            self._set_ai_running(False)

    def soc_gen_kql(self):
        if not self._soc_precheck(): return
        user_input = self.soc_prompt_text.get("1.0", "end").strip()
        threading.Thread(target=self._soc_kql_only_thread, args=(user_input,), daemon=True).start()

    def _soc_kql_only_thread(self, user_input):
        try:
            provider = self._get_provider()
            api_key = self._get_api_key()
            model = self.soc_model_var.get()
            active_hints = self.get_active_hints()
            focus = self.active_flag_var.get()

            # Get Accumulated Report Context
            incident_ctx = self._get_incident_context_for_kql()

            ctx = get_query_context(provider, api_key, user_input, model, self.soc_memory, active_hints, focus, incident_ctx)
            # Update last KQL for self-healing manual runs
            self.last_kql = ctx.get('kql_query')
            self.soc_print(f"\n{Fore.YELLOW}KQL Preview:\n{ctx.get('kql_query')}{Fore.RESET}")
        except Exception as e:
            self.soc_print(str(e))

    def soc_open_manual_kql(self):
        if not self._soc_precheck(): return
        win = tk.Toplevel(self.root)
        win.title("Manual KQL Editor")
        win.geometry("800x600")
        lbl = ttk.Label(win, text="Enter your KQL Query below (Press Run to execute):")
        lbl.pack(padx=10, pady=5, anchor="w")
        text_area = scrolledtext.ScrolledText(win, font=("Consolas", 12))
        text_area.pack(fill="both", expand=True, padx=10, pady=5)
        btn_frame = ttk.Frame(win)
        btn_frame.pack(fill="x", padx=10, pady=10)

        def run_manual():
            query = text_area.get("1.0", "end").strip()
            if not query: return
            win.destroy()
            self._ai_cancel_event.clear()
            self._set_ai_running(True)
            threading.Thread(target=self._soc_manual_thread, args=(query,), daemon=True).start()

        ttk.Button(btn_frame, text="RUN QUERY", command=run_manual).pack(side="right")

    def _soc_manual_thread(self, kql):
        self.soc_print(f"Running Manual KQL: {kql}")
        self.last_kql = kql
        try:
            law_client = LogsQueryClient(credential=DefaultAzureCredential())
            results = execute_kql(law_client, self.workspace_id_var.get(), kql, hours=8760)
            self.soc_last_records = results

            if self._ai_cancelled():
                self.soc_print(f"{Fore.YELLOW}--- MANUAL KQL STOPPED ---{Fore.RESET}")
                return

            if not results:
                self.last_error = "0 records found."
                self.soc_print(f"{Fore.GREEN}0 records found.{Fore.RESET}")
                return

            # FIX: Infer table name correctly to avoid split error on empty string
            table_name = "Unknown"
            if kql:
                # If starts with let, try to find the real table name after last semicolon
                if kql.strip().lower().startswith("let"):
                    parts = kql.strip().split(";")
                    if parts:
                        table_name = parts[-1].strip().split()[0]
                else:
                    table_name = kql.split()[0]

            model = self.soc_model_var.get()
            self._process_results_in_batches(results, table_name, model)

        except Exception as e:
            self.last_error = str(e)
            self.soc_print(f"{Fore.RED}Query Failed: {e}{Fore.RESET}")
        finally:
            self._set_ai_running(False)

    def soc_self_heal(self):
        if not self._soc_precheck(): return
        if not self.last_kql:
            messagebox.showinfo("Info", "No recent KQL query to fix.")
            return

        self.soc_print(f"{Fore.CYAN}--- STARTING SELF-HEAL ---{Fore.RESET}")
        self.soc_print(f"Attempting to fix: {self.last_kql}")
        self.soc_print(f"Last Error: {self.last_error}")

        threading.Thread(target=self._soc_self_heal_thread, daemon=True).start()

    def _soc_self_heal_thread(self):
        try:
            prompt = f"""
            You are a KQL Expert. The following query failed or returned no results.

            Failed Query:
            {self.last_kql}

            Error Message / Issue:
            {self.last_error}

            Task: Fix the KQL syntax, table names, or logic.
            - If "0 records found", maybe the time range is too short (default to 30d) or the 'where' clause is too strict.
            - If syntax error, correct it.

            Return JSON: {{ "fixed_kql": "...", "explanation": "..." }}
            """

            content = ai_chat_completion(
                self._get_provider(), self._get_api_key(), self.soc_model_var.get(),
                [{"role": "user", "content": prompt}], json_mode=True
            )
            data = json.loads(content)
            fixed_kql = data.get("fixed_kql", "")
            explanation = data.get("explanation", "")

            self.soc_print(f"{Fore.YELLOW}Fixed KQL proposed:\n{fixed_kql}{Fore.RESET}")
            self.soc_print(f"Reason: {explanation}")

            # Interactive Confirmation Step (Main Thread)
            self.soc_print("Prompting user to review fix...")
            self.root.after(0, lambda: self._confirm_and_run_fix(fixed_kql))

        except Exception as e:
            self.soc_print(f"{Fore.RED}Self-Heal Failed: {e}{Fore.RESET}")

    def _confirm_and_run_fix(self, fixed_kql):
        # Open a dialog to confirm/edit
        win = tk.Toplevel(self.root)
        win.title("Confirm Fixed KQL")
        win.geometry("800x600")

        lbl = ttk.Label(win, text="The AI suggested the following fix. Review/Edit and Run:")
        lbl.pack(padx=10, pady=5, anchor="w")

        text_area = scrolledtext.ScrolledText(win, font=("Consolas", 12))
        text_area.pack(fill="both", expand=True, padx=10, pady=5)
        text_area.insert("1.0", fixed_kql)

        btn_frame = ttk.Frame(win)
        btn_frame.pack(fill="x", padx=10, pady=10)

        def run():
            query = text_area.get("1.0", "end").strip()
            win.destroy()
            if query:
                self._ai_cancel_event.clear()
                self._set_ai_running(True)
                threading.Thread(target=self._soc_manual_thread, args=(query,), daemon=True).start()

        ttk.Button(btn_frame, text="RUN FIXED QUERY", command=run).pack(side="right")
        ttk.Button(btn_frame, text="Cancel", command=win.destroy).pack(side="right", padx=5)

    def soc_export_log(self):
        f = filedialog.asksaveasfilename(defaultextension=".txt", filetypes=[("Text File", "*.txt")])
        if f:
            if self.logger.save_to_file(f):
                self.soc_print(f"Session saved to {f}")
                if messagebox.askyesno("Integration", "Load this log into Threat Hunter Tab?"):
                    self.th_files.append(f)
                    self.file_listbox.insert("end", os.path.basename(f))
                    self.notebook.select(self.tab_hunter)
            else:
                messagebox.showerror("Error", "Failed to save file.")

    # -------------------------------------------------------------------------
    # SUMMARY TAB (NEW)
    # -------------------------------------------------------------------------
    def _setup_summary_tab(self):
        frame = ttk.LabelFrame(self.tab_summary, text="Confirmed Flags & Notes", padding=10)
        frame.pack(fill="both", expand=True, padx=10, pady=10)

        self.summary_text = scrolledtext.ScrolledText(frame, font=("Consolas", 11), state="disabled")
        self.summary_text.pack(fill="both", expand=True)

        btn_box = ttk.Frame(frame)
        btn_box.pack(fill="x", pady=5)
        ttk.Button(btn_box, text="Refresh Display", command=self._update_summary_display).pack(side="right")

    def _update_summary_display(self):
        self.summary_text.config(state="normal")
        self.summary_text.delete("1.0", "end")

        txt = "=== INVESTIGATION FINDINGS SUMMARY ===\n\n"

        # Sort by Title to keep tidy
        sorted_flags = sorted(self.verified_flags_data, key=lambda x: x.get('title', ''))

        for idx, item in enumerate(sorted_flags, 1):
            txt += f"{idx}. FLAG: {item['title']}\n"
            txt += f"   FOCUS: {item.get('focus_id', 'General')}\n"
            txt += f"   NOTE:  {item['note']}\n"
            txt += "-"*50 + "\n"

        self.summary_text.insert("1.0", txt)
        self.summary_text.config(state="disabled")

    # -------------------------------------------------------------------------
    # INCIDENT REPORT
    # -------------------------------------------------------------------------
    def _setup_incident_tab(self):
        input_frame = ttk.LabelFrame(self.tab_incident, text="Input Data", padding=10)
        input_frame.pack(fill="x", padx=10, pady=5)

        ttk.Label(input_frame, text="Source Logs (optional - .docx, .txt, .pdf):").grid(row=0, column=0, sticky="w")
        ttk.Entry(input_frame, textvariable=self.ir_file_var, width=60).grid(row=0, column=1, padx=5)
        ttk.Button(input_frame, text="Browse", command=self.ir_browse_file).grid(row=0, column=2, padx=5)

        ttk.Label(input_frame, text="Model:").grid(row=0, column=3, sticky="e")
        ir_model_combo = ttk.Combobox(input_frame, textvariable=self.ir_model_var, values=self._get_models_for_provider(), state="readonly", width=25)
        ir_model_combo.grid(row=0, column=4, padx=5)
        self._model_combos.append((self.ir_model_var, ir_model_combo))

        # Internal data sources panel
        src_frame = ttk.LabelFrame(self.tab_incident, text="Internal Data Sources (auto-included)", padding=10)
        src_frame.pack(fill="x", padx=10, pady=5)

        self.ir_src_flags = tk.BooleanVar(value=True)
        self.ir_src_hints = tk.BooleanVar(value=True)
        self.ir_src_bank = tk.BooleanVar(value=True)
        self.ir_src_soc = tk.BooleanVar(value=True)

        ttk.Checkbutton(src_frame, text="Verified Flags & Answers", variable=self.ir_src_flags).grid(row=0, column=0, sticky="w", padx=10)
        ttk.Checkbutton(src_frame, text="CTF Hints & Clues", variable=self.ir_src_hints).grid(row=0, column=1, sticky="w", padx=10)
        ttk.Checkbutton(src_frame, text="Flag Bank Narrative", variable=self.ir_src_bank).grid(row=0, column=2, sticky="w", padx=10)
        ttk.Checkbutton(src_frame, text="SOC Agent Console & Queries", variable=self.ir_src_soc).grid(row=0, column=3, sticky="w", padx=10)

        self.ir_src_status = ttk.Label(src_frame, text="", foreground="gray")
        self.ir_src_status.grid(row=1, column=0, columnspan=4, sticky="w", padx=10, pady=(5,0))

        tpl_frame = ttk.LabelFrame(self.tab_incident, text="Report Template (Editable)", padding=10)
        tpl_frame.pack(fill="both", expand=True, padx=10, pady=5)

        self.ir_template_text = scrolledtext.ScrolledText(tpl_frame, height=10, font=("Consolas", 10))
        self.ir_template_text.pack(fill="both", expand=True)
        self.ir_template_text.insert("1.0", self.ir_template_default)

        out_frame = ttk.LabelFrame(self.tab_incident, text="Generated Report", padding=10)
        out_frame.pack(fill="both", expand=True, padx=10, pady=5)

        self.ir_output_text = scrolledtext.ScrolledText(out_frame, height=15, font=("Consolas", 10))
        self.ir_output_text.pack(fill="both", expand=True)

        btn_box = ttk.Frame(out_frame)
        btn_box.pack(fill="x", pady=5)
        ttk.Button(btn_box, text="▶ GENERATE REPORT", command=self.ir_generate).pack(side="left", padx=5)
        ttk.Button(btn_box, text="💾 Save Report to .txt", command=self.ir_save).pack(side="right", padx=5)

    def ir_browse_file(self):
        f = filedialog.askopenfilename(filetypes=[("Log Files", "*.docx *.txt *.pdf"), ("All Files", "*.*")])
        if f: self.ir_file_var.set(f)

    def _gather_internal_data(self):
        """Collect all available internal data for the incident report."""
        sections = []
        source_counts = {"flags": 0, "hints": 0, "bank": False, "soc_queries": 0, "soc_logs": 0}

        # 1. Verified Flags & Answers
        if self.ir_src_flags.get() and self.verified_flags_data:
            flags_text = "CONFIRMED FINDINGS / FLAG ANSWERS:\n"
            for f in self.verified_flags_data:
                flags_text += f"  - {f.get('title', 'Unknown')}"
                if f.get('note'):
                    flags_text += f" | Answer: {f['note']}"
                if f.get('focus_id') and f['focus_id'] != 'General/All':
                    flags_text += f" [{f['focus_id']}]"
                flags_text += "\n"
            sections.append(flags_text)
            source_counts["flags"] = len(self.verified_flags_data)

        # 2. CTF Hints & AI-generated Clues
        if self.ir_src_hints.get() and self.ctf_hints:
            hints_text = "INVESTIGATION HINTS & AI CLUES:\n"
            for h in self.ctf_hints:
                hints_text += f"  - [{h.get('id', 'General')}] Hint: {h.get('hint', '')}\n"
                if h.get('clue'):
                    hints_text += f"    Clue: {h['clue']}\n"
                if h.get('kql'):
                    hints_text += f"    Suggested KQL: {h['kql']}\n"
            sections.append(hints_text)
            source_counts["hints"] = len(self.ctf_hints)

        # 3. Flag Bank Narrative
        if self.ir_src_bank.get() and self._flag_bank_cache:
            sections.append(f"CURRENT INCIDENT UNDERSTANDING (Flag Bank Narrative):\n{self._flag_bank_cache}")
            source_counts["bank"] = True

        # 4. SOC Agent - Query History & Console Output
        if self.ir_src_soc.get():
            if self.soc_memory:
                soc_text = "SOC AGENT INVESTIGATION QUERIES:\n"
                for m in self.soc_memory:
                    soc_text += f"  - Question: {m.get('user_input', '')}\n"
                    soc_text += f"    KQL: {m.get('kql_query', '')}\n"
                sections.append(soc_text)
                source_counts["soc_queries"] = len(self.soc_memory)

            # Include last query results (up to 5000 chars)
            if self.soc_last_records:
                try:
                    records_text = json.dumps(self.soc_last_records[:30], default=str, indent=1)
                    if len(records_text) > 5000:
                        records_text = records_text[:5000] + "\n... (truncated)"
                    sections.append(f"LATEST SOC QUERY RESULTS ({len(self.soc_last_records)} total records, showing first 30):\n{records_text}")
                    source_counts["soc_logs"] = len(self.soc_last_records)
                except Exception:
                    pass

        return "\n\n".join(sections), source_counts

    def ir_generate(self):
        if not self._get_api_key():
            messagebox.showerror("Error", "Missing API Key in Configuration tab.")
            return

        filepath = self.ir_file_var.get()
        internal_data, counts = self._gather_internal_data()

        # Check we have at least some data to work with
        has_file = filepath and os.path.exists(filepath)
        has_internal = bool(internal_data.strip())

        if not has_file and not has_internal:
            messagebox.showerror("Error", "No data available. Either select a source file, or generate some findings using the other tabs first (SOC Agent, Threat Hunter, Flag Hints).")
            return

        # Show what data sources are being used
        source_parts = []
        if has_file:
            source_parts.append(f"File: {os.path.basename(filepath)}")
        if counts["flags"]:
            source_parts.append(f"{counts['flags']} verified flag(s)")
        if counts["hints"]:
            source_parts.append(f"{counts['hints']} hint(s)")
        if counts["bank"]:
            source_parts.append("Flag Bank narrative")
        if counts["soc_queries"]:
            source_parts.append(f"{counts['soc_queries']} SOC query/queries")
        if counts["soc_logs"]:
            source_parts.append(f"{counts['soc_logs']} log record(s)")

        status_msg = "Sources: " + ", ".join(source_parts)
        self.ir_src_status.config(text=status_msg, foreground="blue")

        self.ir_output_text.delete("1.0", "end")
        self.ir_output_text.insert("1.0", f"Generating report using: {', '.join(source_parts)}...\nPlease wait.\n")

        template_content = self.ir_template_text.get("1.0", "end").strip()

        threading.Thread(target=self._ir_process_thread, args=(filepath if has_file else None, template_content, internal_data), daemon=True).start()

    def auto_generate_incident_report(self):
        """Automatically runs the report generation with current data."""
        if not self._get_api_key():
            return

        filepath = self.ir_file_var.get()

        # Fallback to first TH file if IR file not set
        if (not filepath or not os.path.exists(filepath)) and self.th_files:
            filepath = self.th_files[0]
            self.ir_file_var.set(filepath)

        has_file = filepath and os.path.exists(filepath)
        internal_data, counts = self._gather_internal_data()
        has_internal = bool(internal_data.strip())

        if not has_file and not has_internal:
            print("Auto-Report skipped: No data available.")
            return

        self.ir_output_text.delete("1.0", "end")
        self.ir_output_text.insert("1.0", "[AUTO-UPDATE] Generating report from all available data...\n")

        template_content = self.ir_template_text.get("1.0", "end").strip()

        threading.Thread(target=self._ir_process_thread, args=(filepath if has_file else None, template_content, internal_data), daemon=True).start()

    def _ir_process_thread(self, filepath, template_content, internal_data=""):
        try:
            # Extract file content if a file is provided
            file_section = ""
            if filepath:
                chunks = extract_text_from_file(filepath)
                log_content = "\n".join(chunks)
                if log_content.strip():
                    file_section = f"""
=== UPLOADED DOCUMENT (from {os.path.basename(filepath)}) ===
{log_content[:20000]}"""

            # Build the full data payload
            data_payload = ""
            if internal_data.strip():
                data_payload += f"""
=== INVESTIGATION DATA (from program) ===
{internal_data}
"""
            if file_section:
                data_payload += file_section

            if not data_payload.strip():
                self._ir_update_output("Error: No usable data found.")
                return

            user_prompt = f"""Using ALL of the investigation data below, generate a complete incident report.

TASK 1: TIMELINE OF EVENTS
  - Create a chronological timeline (YYYY-MM-DD HH:MM:SS UTC format).
  - Use timestamps from logs, alerts, and findings.
  - If exact timestamps are unavailable, note the sequence of events.

TASK 2: INVESTIGATION REPORT
  - Fill in the template below using ALL available data.
  - Include the specific flag answers/artifacts found.
  - Reference the evidence and KQL queries used.
  - Provide actionable recommendations.

{data_payload}

=== REPORT TEMPLATE ===
{template_content}
"""
            report = ai_chat_completion(
                self._get_provider(), self._get_api_key(), self.ir_model_var.get(),
                [{"role": "system", "content": SYSTEM_PROMPT_INCIDENT_REPORT}, {"role": "user", "content": user_prompt}],
                temperature=0.3
            )
            self._ir_update_output(report, clear=True)
        except Exception as e:
            self._ir_update_output(f"Error during generation: {e}")

    def _ir_update_output(self, text, clear=False):
        def _update():
            if clear: self.ir_output_text.delete("1.0", "end")
            self.ir_output_text.insert("end", text)
        self.root.after(0, _update)

    def ir_save(self):
        content = self.ir_output_text.get("1.0", "end").strip()
        if not content: return
        f = filedialog.asksaveasfilename(defaultextension=".txt", filetypes=[("Text File", "*.txt")])
        if f:
            try:
                with open(f, "w", encoding="utf-8") as file: file.write(content)
                messagebox.showinfo("Saved", "Report saved successfully.")
            except Exception as e:
                messagebox.showerror("Error", str(e))

    # -------------------------------------------------------------------------
    # SESSION MANAGER
    # -------------------------------------------------------------------------
    def _setup_session_tab(self):
        frame = ttk.LabelFrame(self.tab_session, text="Session Management", padding=20)
        frame.pack(fill="both", expand=True, padx=20, pady=20)

        lbl = ttk.Label(frame, text="Save or Load the entire application state (including API Keys).", font=("Segoe UI", 12))
        lbl.pack(pady=10)

        btn_save = ttk.Button(frame, text="💾 Save Full Session", command=self.save_full_session)
        btn_save.pack(fill="x", pady=5)

        btn_load = ttk.Button(frame, text="📂 Load Full Session", command=self.load_full_session)
        btn_load.pack(fill="x", pady=5)

        self.session_status = ttk.Label(frame, text="Ready", foreground="gray")
        self.session_status.pack(pady=20)

    # -------------------------------------------------------------------------
    # HOW-TO GUIDE
    # -------------------------------------------------------------------------
    def _setup_guide_tab(self):
        guide_text = scrolledtext.ScrolledText(self.tab_guide, font=("Consolas", 11), wrap="word")
        guide_text.pack(fill="both", expand=True, padx=10, pady=10)
        guide_content = """\
================================================================================
        UNIFIED SOC ANALYST & THREAT HUNTER - HOW-TO GUIDE
================================================================================

TABLE OF CONTENTS
-----------------
  1. Getting Started (Prerequisites & Configuration)
  2. Tab-by-Tab Walkthrough
     2a. Configuration
     2b. Flag Hints
     2c. Threat Hunter
     2d. Azure SOC Agent
     2e. Flag Bank (Context)
     2f. Flag Summary
     2g. Report Editor
     2h. Incident Report Generator
     2i. Session Manager
  3. Typical Investigation Workflow
  4. Tips & Troubleshooting


================================================================================
1. GETTING STARTED
================================================================================

PREREQUISITES
  - Python 3.9+
  - Required libraries (install with pip):
      pip install openai pypdf azure-identity azure-monitor-query pandas
      pip install colorama python-docx
  - Optional AI providers (install as needed):
      pip install anthropic          (for Claude support)
      pip install google-generativeai (for Gemini support)
  - An API key for your chosen AI provider:
      OpenAI: https://platform.openai.com
      Google Gemini: https://aistudio.google.com
      Anthropic Claude: https://console.anthropic.com
  - An Azure Log Analytics Workspace ID (for the SOC Agent tab)
  - Azure credentials configured (az login, or environment variables)

FIRST LAUNCH
  1. Run the tool:  python unifiedsoctool.py
  2. Go to the "Configuration" tab.
  3. Select your AI Provider (OpenAI, Gemini, or Claude).
  4. Paste the corresponding API key.
  5. Paste your Log Analytics Workspace ID (for SOC Agent tab).
  6. Click "Save / Validate".
  7. You are now ready to use all features.


================================================================================
2. TAB-BY-TAB WALKTHROUGH
================================================================================

--------------------------------------------------------------------------------
2a. CONFIGURATION TAB
--------------------------------------------------------------------------------
  PURPOSE: Store your AI provider settings and API keys in memory.

  HOW TO USE:
  - Select your AI Provider (OpenAI, Gemini, or Claude).
  - Paste your API key(s) in the corresponding field(s) (masked for security).
  - Paste your Azure Log Analytics Workspace ID in the Azure Settings section.
  - Click "Save / Validate" to confirm settings are loaded.
  - These credentials are NOT saved to disk unless you use Session Manager.
  - Tip: API keys can be set via environment variables:
    OPENAI_API_KEY, GEMINI_API_KEY, ANTHROPIC_API_KEY


--------------------------------------------------------------------------------
2b. FLAG HINTS TAB
--------------------------------------------------------------------------------
  PURPOSE: Add and analyze CTF hints. The AI generates investigative clues
  and suggests KQL queries based on each hint.

  HOW TO USE:
  1. Select a flag number from the "Assign to Flag" dropdown (Flag 1-100).
  2. Type or paste the hint text in the text area.
  3. Choose an AI model (gpt-4o is recommended for best results).
  4. Click "Analyze Hint".
  5. The AI will return:
     - A clue (what artifact to look for).
     - A suggested KQL query to find it.
  6. Results appear in the "Active Hints" panel below.

  TIP: Use the "Investigation Focus" dropdown at the top of the window to
  filter which hints are sent to the AI during investigations. Set it to
  a specific flag to focus, or "General/All" to use all hints.


--------------------------------------------------------------------------------
2c. THREAT HUNTER TAB
--------------------------------------------------------------------------------
  PURPOSE: Load files (PDF, TXT, LOG, JSONL) and let the AI scan them for
  security threats, flags, and suspicious activity.

  HOW TO USE:
  1. Click "Add Files" and select one or more supported files.
  2. Choose an AI model.
  3. Click "START HUNT".
  4. The AI reads the files in batches and looks for threats.
  5. When a finding is discovered, it appears in the "Potential Flag" editor:
     - Click "Verify & Save" to confirm and save the finding.
     - Click "Discard" to skip it.
  6. Verified findings go to Flag Summary and trigger report updates.

  MANUAL FLAG: Click "Manually Add Flag" to enter a finding by hand
  (useful if you spot something the AI missed).


--------------------------------------------------------------------------------
2d. AZURE SOC AGENT TAB
--------------------------------------------------------------------------------
  PURPOSE: Query Azure Log Analytics using natural language. The AI converts
  your question into KQL and executes it against your workspace.

  HOW TO USE:
  1. Type your investigation question in the prompt box, e.g.:
     "Find all failed sign-ins from external IPs in the last 7 days"
     "Show process executions on device WKS01"
     "Search for PowerShell activity across all devices"
  2. Click "Run AI Investigation":
     - AI generates a KQL query.
     - Query runs against Azure Log Analytics (up to 1 year of data).
     - Results are analyzed for threats automatically.
     - Findings are pushed to the Threat Hunter for verification.
  3. Click "Stop AI" at any time to halt a running investigation.

  OTHER BUTTONS:
  - "Generate KQL Only": See the AI-generated query without running it.
  - "Open Manual KQL Editor": Write and execute your own KQL query directly.
  - "Export Log to File": Save the console output to a .txt file.
  - "Self-Heal Last KQL": If a query fails, the AI attempts to fix and
    re-run it automatically.

  TIME RANGE:
  - The tool searches up to 1 year (365 days) of log data by default.
  - The AI can adjust this based on your query (e.g., "last 7 days").
  - Manual KQL queries also support up to 1 year.

  TABLES:
  You can query ANY table in your Azure Log Analytics workspace.
  Just name the table in your query and the AI will use it.
  Common tables include:
  - DeviceProcessEvents     - DeviceNetworkEvents
  - DeviceLogonEvents       - DeviceFileEvents
  - DeviceRegistryEvents    - DeviceImageLoadEvents
  - EmailEvents             - EmailAttachmentInfo
  - IdentityLogonEvents     - CloudAppEvents
  - SigninLogs              - AuditLogs
  - AzureActivity           - SecurityAlert
  - SecurityIncident        - SecurityEvent
  - Syslog                  - CommonSecurityLog
  - ThreatIntelligenceIndicator  - BehaviorAnalytics
  - ...and any other table in your workspace


--------------------------------------------------------------------------------
2e. FLAG BANK (CONTEXT) TAB
--------------------------------------------------------------------------------
  PURPOSE: AI-generated narrative that connects all your verified findings
  into a coherent story.

  HOW TO USE:
  1. Click "Update Understanding with AI" to generate/refresh the narrative.
  2. The AI reviews all verified flags and hints, then writes a cohesive
     summary connecting the dots.
  3. Click "Export Narrative to Incident Report" to push this context into
     the Incident Report Generator's template.


--------------------------------------------------------------------------------
2f. FLAG SUMMARY TAB
--------------------------------------------------------------------------------
  PURPOSE: View all verified flags and their analyst notes in one place.

  HOW TO USE:
  - Click "Refresh Display" to update.
  - Flags are listed with their title, focus ID, and note.
  - This is your master list of confirmed findings.


--------------------------------------------------------------------------------
2g. REPORT EDITOR TAB (Threat Hunt Reporter)
--------------------------------------------------------------------------------
  PURPOSE: Build a formal threat hunt report with findings, methodology,
  KQL queries, and screenshots, then export as a Word (.docx) document.

  HOW TO USE:
  1. Fill in "Author Name" and "Event/Engagement".
  2. For each finding:
     a. Enter Challenge/Title, select a Category.
     b. Describe the finding and your methodology.
     c. Paste the KQL query used.
     d. Enter the flag/artifact value.
     e. Optionally attach screenshots with captions.
     f. Click "Add Entry to Report".
  3. Review entries in the "Entries Queued" list.
  4. Click "GENERATE WORD DOC REPORT" to create the .docx file.

  SUB-TAB "Help & Instructions": Quick reference for this module.


--------------------------------------------------------------------------------
2h. INCIDENT REPORT GENERATOR TAB
--------------------------------------------------------------------------------
  PURPOSE: Generate a formal incident investigation report using ALL data
  available in the program, plus optional uploaded documents.

  DATA SOURCES (auto-included via checkboxes):
  - Verified Flags & Answers: All confirmed findings with exact answers.
  - CTF Hints & Clues: All hints and AI-generated investigation clues.
  - Flag Bank Narrative: The AI-generated incident understanding story.
  - SOC Agent Console & Queries: All KQL queries run, plus latest results.
  - Uploaded Document (optional): Any .docx, .txt, or .pdf file.

  HOW TO USE:
  1. (Optional) Click "Browse" to attach a source log file.
  2. Check/uncheck the internal data sources you want to include.
  3. Edit the report template if desired (Who/What/When/Where/Why/How).
  4. Click "Generate Report":
     - The AI gathers all checked internal data automatically.
     - Reads the uploaded document (if any).
     - Creates a timeline of events.
     - Fills in the template with all findings and flag answers.
     - References KQL queries and evidence used.
  5. Click "Save Report to .txt" to export.

  NOTE: You do NOT need an uploaded file. The report can be generated
  entirely from internal data (flags, hints, SOC queries, etc.).


--------------------------------------------------------------------------------
2i. SESSION MANAGER TAB
--------------------------------------------------------------------------------
  PURPOSE: Save and restore your entire application state.

  HOW TO USE:
  - "Save Full Session": Saves everything (API keys, hints, flags, queries,
    reports, console output) to a JSON file.
  - "Load Full Session": Restores a previously saved session.

  TIP: Save your session frequently, especially during long investigations.


================================================================================
3. TYPICAL INVESTIGATION WORKFLOW
================================================================================

  Step 1: CONFIGURE
     Set up API key and Workspace ID in the Configuration tab.

  Step 2: ADD HINTS (if applicable)
     Go to Flag Hints tab and add any hints or leads you have.
     Let the AI analyze them for investigative clues.

  Step 3: INVESTIGATE WITH SOC AGENT
     Go to the Azure SOC Agent tab.
     Ask questions in natural language:
       "Show all sign-in failures for user john@company.com"
       "Find devices with suspicious PowerShell execution"
       "Search for lateral movement activity"
     The AI generates KQL, runs it, and analyzes results.

  Step 4: REVIEW & VERIFY FINDINGS
     Switch to Threat Hunter when findings appear.
     Verify or discard each finding.
     Add manual flags for anything the AI missed.

  Step 5: BUILD CONTEXT
     Go to Flag Bank and click "Update Understanding with AI".
     Review the AI's narrative connecting all findings.

  Step 6: GENERATE REPORTS
     Use the Report Editor for detailed per-finding reports with screenshots.
     Use the Incident Report Generator for a formal investigation summary.

  Step 7: SAVE SESSION
     Save your session in the Session Manager tab for future reference.


================================================================================
4. TIPS & TROUBLESHOOTING
================================================================================

  QUERY TIPS:
  - Be specific in your questions: "failed logins for admin on WKS01"
    works better than "show me everything suspicious".
  - If a query fails, click "Self-Heal Last KQL" to auto-fix.
  - Use "Generate KQL Only" to preview queries before running them.
  - For complex queries the AI can't handle, use "Open Manual KQL Editor".

  KQL SYNTAX QUICK REFERENCE:
  - Filter:    | where ColumnName == "value"
  - Contains:  | where ColumnName contains "partial"
  - Starts:    | where ColumnName startswith "prefix"
  - Time:      | where TimeGenerated > ago(7d)
  - Limit:     | take 100
  - Sort:      | order by TimeGenerated desc
  - Search:    search "keyword" | take 100

  COMMON ISSUES:
  - "No results found": Try a wider time range or broader search terms.
    The tool searches up to 1 year by default.
  - "Query failed": Click "Self-Heal Last KQL". The AI will diagnose and
    fix the syntax error.
  - "API key error": Make sure your API key is valid and has credits.
    Ensure the correct provider is selected in the Configuration tab.
  - "Azure auth error": Run "az login" in your terminal first, or set
    Azure credential environment variables.
  - "Missing library": Run the pip install command shown at startup.

  KEYBOARD & MOUSE:
  - Mousewheel scrolls all text areas.
  - Standard copy/paste shortcuts work everywhere (Ctrl+C, Ctrl+V).
  - Tab key navigates between fields.

================================================================================
"""
        guide_text.insert("1.0", guide_content)
        guide_text.config(state="disabled")

    def save_full_session(self):
        data = {
            "config": {
                "provider": self.provider_var.get(),
                "api_key_openai": self.api_key_vars["OpenAI"].get(),
                "api_key_gemini": self.api_key_vars["Gemini"].get(),
                "api_key_claude": self.api_key_vars["Claude"].get(),
                "workspace_id": self.workspace_id_var.get(),
                "custom_models": self.custom_models
            },
            "hints": {
                "ctf_hints": self.ctf_hints,
                "model": self.hint_model_var.get(),
                "display_text": self.hint_display.get("1.0", "end")
            },
            "hunter": {
                "files": self.th_files,
                "found_flags": list(self.th_found_flags),
                "model": self.th_model_var.get(),
                "history_text": self.th_history.get("1.0", "end"),
                "verified_data": self.verified_flags_data # Save verified data
            },
            "soc": {
                "memory": self.soc_memory,
                "model": self.soc_model_var.get(),
                "console_text": self.soc_console.get("1.0", "end")
            },
            "reporter": self.reporter.get_state(),
            "incident": {
                "file": self.ir_file_var.get(),
                "model": self.ir_model_var.get(),
                "template": self.ir_template_text.get("1.0", "end"),
                "output": self.ir_output_text.get("1.0", "end")
            },
            "flag_bank": {
                "text": self.flag_bank_text.get("1.0", "end")
            }
        }
        f = filedialog.asksaveasfilename(defaultextension=".json", filetypes=[("JSON Session", "*.json")])
        if f:
            try:
                with open(f, 'w') as file: json.dump(data, file, indent=4)
                self.session_status.config(text=f"Saved to {os.path.basename(f)}", foreground="green")
            except Exception as e:
                messagebox.showerror("Save Error", str(e))

    def load_full_session(self):
        f = filedialog.askopenfilename(filetypes=[("JSON Session", "*.json")])
        if not f: return
        try:
            with open(f, 'r') as file: data = json.load(file)

            if "config" in data:
                cfg = data["config"]
                self.provider_var.set(cfg.get("provider", "OpenAI"))
                # Support both old single-key and new multi-key formats
                if "api_key_openai" in cfg:
                    self.api_key_vars["OpenAI"].set(cfg.get("api_key_openai", ""))
                    self.api_key_vars["Gemini"].set(cfg.get("api_key_gemini", ""))
                    self.api_key_vars["Claude"].set(cfg.get("api_key_claude", ""))
                elif "api_key" in cfg:
                    # Legacy format: single OpenAI key
                    self.api_key_vars["OpenAI"].set(cfg.get("api_key", ""))
                self.workspace_id_var.set(cfg.get("workspace_id", ""))
                # Restore custom models before triggering provider change
                saved_custom = cfg.get("custom_models", {})
                for provider in self.custom_models:
                    self.custom_models[provider] = saved_custom.get(provider, [])
                self._on_provider_changed()

            if "hints" in data:
                self.ctf_hints = data["hints"].get("ctf_hints", [])
                self.hint_model_var.set(data["hints"].get("model", DEFAULT_MODEL))
                self.hint_display.config(state="normal")
                self.hint_display.delete("1.0", "end")
                self.hint_display.insert("1.0", data["hints"].get("display_text", ""))
                self.hint_display.config(state="disabled")

            if "hunter" in data:
                self.th_files = data["hunter"].get("files", [])
                self.file_listbox.delete(0, "end")
                for file_path in self.th_files: self.file_listbox.insert("end", os.path.basename(file_path))
                self.th_found_flags = set(data["hunter"].get("found_flags", []))
                self.verified_flags_data = data["hunter"].get("verified_data", []) # Load verified data
                self.th_model_var.set(data["hunter"].get("model", DEFAULT_MODEL))
                self.th_history.config(state="normal")
                self.th_history.delete("1.0", "end")
                self.th_history.insert("1.0", data["hunter"].get("history_text", ""))
                self.th_history.config(state="disabled")
                self._update_summary_display() # Update Summary Tab

            if "soc" in data:
                self.soc_memory = data["soc"].get("memory", [])
                self.soc_model_var.set(data["soc"].get("model", DEFAULT_MODEL))
                self.soc_console.config(state="normal")
                self.soc_console.delete("1.0", "end")
                self.soc_console.insert("1.0", data["soc"].get("console_text", ""))
                self.soc_console.config(state="disabled")

            if "reporter" in data: self.reporter.set_state(data["reporter"])

            if "incident" in data:
                self.ir_file_var.set(data["incident"].get("file", ""))
                self.ir_model_var.set(data["incident"].get("model", DEFAULT_MODEL))
                self.ir_template_text.delete("1.0", "end")
                self.ir_template_text.insert("1.0", data["incident"].get("template", ""))
                self.ir_output_text.delete("1.0", "end")
                self.ir_output_text.insert("1.0", data["incident"].get("output", ""))

            if "flag_bank" in data:
                bank_text = data["flag_bank"].get("text", "")
                self.flag_bank_text.delete("1.0", "end")
                self.flag_bank_text.insert("1.0", bank_text)
                self._flag_bank_cache = bank_text

            self.session_status.config(text=f"Loaded {os.path.basename(f)}", foreground="green")
            messagebox.showinfo("Session Loaded", "Full session state restored.")
        except Exception as e:
            messagebox.showerror("Load Error", str(e))

if __name__ == "__main__":
    init(autoreset=True)
    root = tk.Tk()
    app = UnifiedSOCTool(root)
    root.mainloop()